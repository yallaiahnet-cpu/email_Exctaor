from flask import Flask, request, jsonify, render_template, send_file
import json
import re
import os
from datetime import datetime
from typing import Optional
from dotenv import load_dotenv
from cleaning_jd import EmailExtractor
from send_email import EmailSender
from llm_exctration import ResumeOptimizer
from document_creation import generate_resume_style_1
from job_scraper import JobScraper

# Load environment variables
load_dotenv()

app = Flask(__name__)

# Initialize EmailSender - load from email.json instead of environment variables
# Default will be the first email in email.json, or empty if file doesn't exist
SMTP_USERNAME = ''
SMTP_PASSWORD = ''
RESUME_FILE = ''

# Try to load first email from email.json as default
try:
    email_json_path = os.path.join(os.getcwd(), 'email.json')
    if os.path.exists(email_json_path):
        with open(email_json_path, 'r') as f:
            emails_data = json.load(f)
            # Get first email entry
            if emails_data and isinstance(emails_data, dict):
                first_email = list(emails_data.keys())[0] if emails_data else ''
                if first_email and first_email in emails_data:
                    entry = emails_data[first_email]
                    if isinstance(entry, list) and len(entry) > 0:
                        first_entry = entry[0]
                        SMTP_USERNAME = first_email
                        SMTP_PASSWORD = first_entry.get('smtp_password') or first_entry.get('app_password') or ''
except Exception as e:
    print(f"⚠️ Could not load default email from email.json: {e}")

email_sender = EmailSender(
    sender_email=SMTP_USERNAME,
    app_password=SMTP_PASSWORD,
    resume_file=RESUME_FILE
)

EMAIL_ARCHIVE_PATH = os.path.join(os.getcwd(), 'sent_emails.json')


def _normalise_recipient_list(value: str) -> list:
    """Convert a comma-separated string of emails into a clean list of strings."""
    if not value:
        return []
    return [item.strip() for item in value.split(',') if item.strip()]


def extract_phone_from_text(text: str) -> Optional[str]:
    """Extract phone number from text using regex patterns, including extensions"""
    if not text:
        return None
    
    # First, try patterns with prefixes (Ph:, Phone:, etc.) - highest priority
    # Match "Ph: 9724402129" or "Phone: 972-440-2129" etc.
    prefix_patterns = [
        # Pattern for "Ph: 9724402129" - captures digits after colon/space
        r'(?:Ph|Phone|Tel|Telephone|Contact|Call|Mobile|Cell)[:\s]+([\d\s\-\(\)\.]+(?:ext|extension|x|#)?[\s:\.]?\d*)',
        # Direct digits after prefix (10-15 digits)
        r'(?:Ph|Phone|Tel|Telephone|Contact|Call|Mobile|Cell)[:\s]+(\d{10,15})',
        # With optional formatting
        r'(?:Ph|Phone|Tel|Telephone|Contact|Call|Mobile|Cell)[:\s]+([+\d\s\-\(\)\.]+)',
    ]
    
    for pattern in prefix_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            for match in matches:
                phone = match.strip() if isinstance(match, str) else str(match).strip()
                # Clean up but preserve structure
                phone = re.sub(r'\s+', ' ', phone)
                # Extract digits for validation (excluding extension)
                phone_main = phone.split('ext')[0].split('extension')[0].split('x')[0].split('#')[0]
                digits_only = re.sub(r'[^\d]', '', phone_main)
                if len(digits_only) >= 10:
                    print(f"✅ Extracted phone with prefix pattern: {phone} (digits: {len(digits_only)})")
                    return phone.strip()
    
    # Then try standard phone formats
    standard_patterns = [
        # US format: (XXX) XXX-XXXX with optional extension
        r'(\(?\d{3}\)?\s?\d{3}[-.\s]?\d{4}(?:\s*(?:ext|extension|x|#)[\s:\.]?\d+)?)',
        # Standard format: XXX-XXX-XXXX or XXX.XXX.XXXX
        r'(\d{3}[-.\s]?\d{3}[-.\s]?\d{4}(?:\s*(?:ext|extension|x|#)[\s:\.]?\d+)?)',
        # 10-digit standalone
        r'(\b\d{10}\b(?:\s*(?:ext|extension|x|#)[\s:\.]?\d+)?)',
        # International format
        r'(\+\d{1,3}[-.\s]?\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,9}(?:\s*(?:ext|extension|x|#)[\s:\.]?\d+)?)',
    ]
    
    for pattern in standard_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            phone = matches[0].strip() if isinstance(matches[0], str) else str(matches[0]).strip()
            phone = re.sub(r'\s+', ' ', phone)
            digits_only = re.sub(r'[^\d]', '', phone.split('ext')[0].split('extension')[0].split('x')[0].split('#')[0])
            if len(digits_only) >= 10:
                print(f"✅ Extracted phone with standard pattern: {phone}")
                return phone.strip()
    
    print(f"⚠️ No phone number found in text")
    return None

def archive_email_metadata(email_payload: dict) -> None:
    """
    Persist the outgoing email payload so we keep a simple JSON archive
    of every email that was attempted.
    """
    try:
        existing_records = []
        if os.path.exists(EMAIL_ARCHIVE_PATH):
            with open(EMAIL_ARCHIVE_PATH, 'r', encoding='utf-8') as archive_file:
                loaded = json.load(archive_file)
                if isinstance(loaded, list):
                    existing_records = loaded

        existing_records.append(email_payload)

        with open(EMAIL_ARCHIVE_PATH, 'w', encoding='utf-8') as archive_file:
            json.dump(existing_records, archive_file, indent=2, ensure_ascii=False)
    except Exception as archive_error:
        print(f"⚠️ Failed to archive email payload: {archive_error}")

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/prompt_generator')
def prompt_generator():
    return render_template('PROMPT_GENARTIOR.HTML')

@app.route('/dotnet_prompt_generator')
def dotnet_prompt_generator():
    return render_template('dotnet_prompt_generator.html')

@app.route('/fullstack_prompt_generator')
def fullstack_prompt_generator():
    return render_template('fullstack_prompt_generator.html')

@app.route('/resume_prompt_generator')
def resume_prompt_generator():
    return render_template('resume_prompt_generator.html')

@app.route('/resume_generator')
def resume_generator():
    return render_template('resume_generator.html')

@app.route('/fulltime')
def fulltime():
    return render_template('fulltime.html')

@app.route('/job_scraper')
def job_scraper():
    return render_template('job_scraper.html')

@app.route('/project_summary_prompt')
def project_summary_prompt():
    return render_template('project_summary_prompt.html')

@app.route('/sent_emails_viewer')
def sent_emails_viewer():
    return render_template('sent_emails_viewer.html')

@app.route('/sent_emails_data', methods=['GET'])
def sent_emails_data():
    """Return sent emails data from JSON archive"""
    try:
        if not os.path.exists(EMAIL_ARCHIVE_PATH):
            return jsonify({'emails': []}), 200
        
        with open(EMAIL_ARCHIVE_PATH, 'r', encoding='utf-8') as f:
            emails = json.load(f)
        
        if not isinstance(emails, list):
            emails = []
        
        return jsonify({'emails': emails}), 200
    except Exception as e:
        print(f"❌ Error reading sent emails: {e}")
        return jsonify({'error': str(e), 'emails': []}), 500

@app.route('/extract_docx_text', methods=['POST'])
def extract_docx_text():
    """Extract text content from DOCX file"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Check if it's a DOCX file
        if not file.filename.lower().endswith(('.docx', '.doc')):
            return jsonify({'error': 'File must be a DOCX or DOC file'}), 400
        
        # Try to extract text using python-docx
        try:
            from docx import Document
            import io
            
            # Read file content
            file_content = file.read()
            doc = Document(io.BytesIO(file_content))
            
            # Extract text from all paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            extracted_text = '\n'.join(text_content)
            
            if not extracted_text.strip():
                return jsonify({'error': 'No text content found in the document'}), 400
            
            return jsonify({'text': extracted_text, 'success': True}), 200
            
        except ImportError:
            return jsonify({'error': 'python-docx library not installed. Please install it using: pip install python-docx'}), 500
        except Exception as e:
            return jsonify({'error': f'Error extracting text: {str(e)}'}), 500
            
    except Exception as e:
        return jsonify({'error': f'Server error: {str(e)}'}), 500

@app.route('/clean_job_description', methods=['POST'])
def clean_job_description():
    try:
        # Get input from frontend
        print("=== Receiving input from frontend ===")
        data = request.get_json()
        raw_text = data.get('raw_text', '')
        sender_email = data.get('sender_email', '')  # Optional sender email
        
        print(f"Input length: {len(raw_text)} chars")
        print(f"📧 Sender email received: '{sender_email}' (empty: {not sender_email})")
        
        # Get years_of_experience from email.json if sender_email is provided
        years_of_experience = "10+ years"  # Default
        if sender_email:
            try:
                email_json_path = os.path.join(os.getcwd(), 'email.json')
                if os.path.exists(email_json_path):
                    with open(email_json_path, 'r') as f:
                        emails_data = json.load(f)
                    print(f"📋 Checking email.json for: {sender_email}")
                    print(f"📋 Available emails: {list(emails_data.keys())}")
                    if sender_email in emails_data and isinstance(emails_data[sender_email], list) and len(emails_data[sender_email]) > 0:
                        entry = emails_data[sender_email][0]
                        years_of_experience = entry.get('years_of_experience', '10+ years')
                        print(f"✅ Using years_of_experience: '{years_of_experience}' for {sender_email}")
                    else:
                        print(f"⚠️ Sender email '{sender_email}' not found in email.json, using default: {years_of_experience}")
                else:
                    print(f"⚠️ email.json not found at {email_json_path}, using default: {years_of_experience}")
            except Exception as e:
                print(f"⚠️ Could not load years_of_experience from email.json: {e}")
                import traceback
                traceback.print_exc()
        else:
            print(f"ℹ️ No sender email provided, using default years_of_experience: {years_of_experience}")
        
        # Pass to cleaning_jd.py (port 5002 uses profile 1)
        print("\n=== Calling cleaning_jd.py EmailExtractor ===")
        extractor = EmailExtractor(raw_text, use_profile_2=False, years_of_experience=years_of_experience)
        print("Extractor created")
        
        result_json = extractor.extract_email_info_from_jd(raw_text)
        print(f"\n=== Response from cleaning_jd.py ===\nLength: {len(result_json)} chars\nFirst 200 chars: {result_json[:200]}")
        
        # Try to parse JSON directly first
        try:
            result_dict = json.loads(result_json)
            print("\n=== Successfully parsed JSON on first try ===")
            return jsonify(result_dict), 200
        except json.JSONDecodeError:
            print("\n=== First parse failed, trying to fix newlines ===")
            # Extract JSON and try to fix it
            json_match = re.search(r'\{.*\}', result_json, re.DOTALL)
            if json_match:
                cleaned_json = json_match.group(0)
                # Try multiple parsing strategies
                for strategy in ['direct', 'replace_all']:
                    try:
                        if strategy == 'direct':
                            result_dict = json.loads(cleaned_json)
                            print("=== Parsed successfully with direct method ===")
                        elif strategy == 'replace_all':
                            fixed_json = cleaned_json.replace('\n', ' ').replace('\r', ' ').replace('\t', ' ')
                            result_dict = json.loads(fixed_json)
                            print("=== Parsed successfully by replacing all newlines with spaces ===")
                        return jsonify(result_dict), 200
                    except json.JSONDecodeError:
                        continue
            print("\n❌ No valid JSON found after all attempts")
            return jsonify({'error': 'No valid JSON found in response'}), 500
    except json.JSONDecodeError as e:
        print(f"\n❌ JSON Error: {str(e)}")
        return jsonify({'error': f'Failed to parse JSON: {str(e)}'}), 500
    except Exception as e:
        print(f"\n❌ Exception: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/list_resumes', methods=['GET'])
def list_resumes():
    """List available resume files"""
    try:
        resumes = []
        
        # Check root directory
        root_dir = os.getcwd()
        for file in os.listdir(root_dir):
            if file.endswith(('.pdf', '.docx', '.doc')):
                full_path = os.path.join(root_dir, file)
                if os.path.isfile(full_path):
                    resumes.append(full_path)
        
        # Check generated_resumes directory
        resume_directory = os.path.join(root_dir, 'generated_resumes')
        if os.path.exists(resume_directory):
            for root, dirs, files in os.walk(resume_directory):
                for file in files:
                    if file.endswith(('.pdf', '.docx', '.doc')):
                        full_path = os.path.join(root, file)
                        if full_path not in resumes:
                            resumes.append(full_path)
        
        resumes.sort()
        return jsonify({'resumes': resumes}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/list_desktop_folders', methods=['GET'])
def list_desktop_folders():
    """List folders from Desktop only"""
    try:
        home_path = os.path.expanduser('~')
        desktop_folders = []
        
        # Only Desktop folders
        desktop_path = os.path.join(home_path, 'Desktop')
        if os.path.exists(desktop_path) and os.path.isdir(desktop_path):
            try:
                for item in os.listdir(desktop_path):
                    item_path = os.path.join(desktop_path, item)
                    if os.path.isdir(item_path):
                        desktop_folders.append({
                            'name': item,
                            'path': item_path,
                            'location': item_path,
                            'parent': 'Desktop'
                        })
            except PermissionError:
                pass
        
        desktop_folders.sort(key=lambda x: x['name'].lower())
        return jsonify({'folders': desktop_folders}), 200
    except Exception as e:
        return jsonify({'error': str(e), 'folders': []}), 500

@app.route('/list_location_folders', methods=['POST'])
def list_location_folders():
    """List folders from a specific location (Documents, Downloads, etc.)"""
    try:
        data = request.get_json()
        location_name = data.get('location', '')
        
        if not location_name:
            return jsonify({'error': 'Location not specified', 'folders': []}), 400
        
        home_path = os.path.expanduser('~')
        folders = []
        
        # Get folders from the specified location
        location_path = os.path.join(home_path, location_name)
        if os.path.exists(location_path) and os.path.isdir(location_path):
            try:
                for item in os.listdir(location_path):
                    item_path = os.path.join(location_path, item)
                    if os.path.isdir(item_path):
                        folders.append({
                            'name': item,
                            'path': item_path,
                            'location': item_path,
                            'parent': location_name
                        })
            except PermissionError:
                pass
        
        folders.sort(key=lambda x: x['name'].lower())
        return jsonify({'folders': folders}), 200
    except Exception as e:
        return jsonify({'error': str(e), 'folders': []}), 500

@app.route('/list_available_locations', methods=['GET'])
def list_available_locations():
    """List all available user directory locations"""
    try:
        home_path = os.path.expanduser('~')
        locations = []
        
        # Common user directories (excluding Desktop)
        common_dirs = ['Documents', 'Downloads', 'Movies', 'Music', 'Pictures', 'Public']
        
        for dir_name in common_dirs:
            dir_path = os.path.join(home_path, dir_name)
            if os.path.exists(dir_path) and os.path.isdir(dir_path):
                locations.append({
                    'name': dir_name,
                    'path': dir_path
                })
        
        locations.sort(key=lambda x: x['name'].lower())
        return jsonify({'locations': locations}), 200
    except Exception as e:
        return jsonify({'error': str(e), 'locations': []}), 500

@app.route('/list_folder_files', methods=['POST'])
def list_folder_files():
    """List both folders and files in a specific directory"""
    try:
        data = request.get_json()
        folder_path = data.get('folder_path', '')
        
        if not folder_path or not os.path.exists(folder_path):
            return jsonify({'error': 'Invalid folder path', 'files': [], 'folders': []}), 400
        
        if not os.path.isdir(folder_path):
            return jsonify({'error': 'Path is not a directory', 'files': [], 'folders': []}), 400
        
        files = []
        folders = []
        
        for item in os.listdir(folder_path):
            item_path = os.path.join(folder_path, item)
            try:
                if os.path.isdir(item_path):
                    folders.append({
                        'name': item,
                        'path': item_path
                    })
                elif os.path.isfile(item_path) and item.endswith(('.pdf', '.docx', '.doc')):
                    files.append({
                        'name': item,
                        'path': item_path
                    })
            except (PermissionError, OSError):
                # Skip items we can't access
                continue
        
        files.sort(key=lambda x: x['name'].lower())
        folders.sort(key=lambda x: x['name'].lower())
        
        return jsonify({
            'files': files,
            'folders': folders,
            'current_path': folder_path
        }), 200
    except Exception as e:
        return jsonify({'error': str(e), 'files': [], 'folders': []}), 500

@app.route('/create_resume', methods=['POST'])
def create_resume():
    """Create optimized resume using ResumeOptimizer and convert to DOCX"""
    try:
        data = request.get_json()
        job_description = data.get('job_description', '')
        
        if not job_description:
            return jsonify({'error': 'Job description is required'}), 400
        
        print(f"\n=== Creating Resume ===")
        print(f"Job description length: {len(job_description)} chars")
        
        # Step 1: Initialize ResumeOptimizer
        optimizer = ResumeOptimizer(job_description)
        
        # Step 2: Extract skills from job description
        print("\n=== Step 1: Extracting skills from JD ===")
        extracted_skills = optimizer.extract_skills()
        print(f"✅ Skills extracted: {len(extracted_skills)} chars")
        
        # Step 3: Generate optimized resume JSON
        print("\n=== Step 2: Generating optimized resume ===")
        optimizer.generate_resume()
        
        if not optimizer.resume_json:
            return jsonify({'error': 'Failed to generate resume JSON'}), 500
        
        # Validate and ensure JSON is a dict, not a string
        print("\n=== Step 3: Validating JSON ===")
        if isinstance(optimizer.resume_json, str):
            optimizer.resume_json = json.loads(optimizer.resume_json)
        
        json.dumps(optimizer.resume_json)  # This will raise error if invalid
        print("✅ JSON is valid")
        
        # Step 4: Save JSON file
        print("\n=== Step 4: Saving JSON file ===")
        output_dir = "resumes"
        os.makedirs(output_dir, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        json_filename = f"optimized_resume_{timestamp}.json"
        json_resume_path = os.path.join(output_dir, json_filename)
        
        with open(json_resume_path, 'w') as f:
            json.dump(optimizer.resume_json, f, indent=2)
        print(f"✅ JSON saved at: {json_resume_path}")
        
        # Step 5: Convert JSON to DOCX using style 5
        print("\n=== Step 5: Converting to DOCX (Style 5) ===")
        resume_directory = "generated_resumes"
        os.makedirs(resume_directory, exist_ok=True)
        
        from document_creation import generate_resume_style_5
        docx_resume_path = generate_resume_style_5(json_resume_path, resume_directory)
        
        if not docx_resume_path or not os.path.exists(docx_resume_path):
            return jsonify({'error': 'Failed to create resume document'}), 500
        
        print(f"✅ Resume DOCX created at: {docx_resume_path}")
        return jsonify({'resume_path': docx_resume_path, 'success': True}), 200
        
    except Exception as e:
        print(f"❌ Error creating resume: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/generate_resume_from_json', methods=['POST'])
def generate_resume_from_json():
    """Generate resume directly from JSON input without LLM"""
    try:
        data = request.get_json()
        json_resume_data = data.get('json_resume_data', '')
        
        if not json_resume_data:
            return jsonify({'error': 'JSON resume data is required'}), 400
        
        # Validate JSON
        try:
            resume_json = json.loads(json_resume_data)
        except json.JSONDecodeError as e:
            return jsonify({'error': f'Invalid JSON: {str(e)}'}), 400
        
        # Save JSON to temp file
        output_dir = "resumes"
        os.makedirs(output_dir, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        json_filename = f"resume_from_json_{timestamp}.json"
        json_resume_path = os.path.join(output_dir, json_filename)
        
        with open(json_resume_path, 'w') as f:
            json.dump(resume_json, f, indent=2)
        print(f"✅ JSON saved at: {json_resume_path}")
        
        # Load keywords from bold_words.json for bolding
        keywords_list = []
        try:
            bold_words_path = os.path.join(os.getcwd(), 'bold_words.json')
            if os.path.exists(bold_words_path):
                with open(bold_words_path, 'r') as f:
                    bold_words_data = json.load(f)
                    if isinstance(bold_words_data, dict) and 'keywords' in bold_words_data:
                        keywords_list = bold_words_data['keywords']
                        print(f"✅ Loaded {len(keywords_list)} keywords from bold_words.json")
                    elif isinstance(bold_words_data, list):
                        keywords_list = bold_words_data
                        print(f"✅ Loaded {len(keywords_list)} keywords from bold_words.json (list format)")
        except Exception as e:
            print(f"⚠️ Could not load keywords from bold_words.json: {e}")
        
        # Generate DOCX using style 5
        print("\n=== Converting JSON to DOCX (Style 5) ===")
        resume_directory = "generated_resumes"
        os.makedirs(resume_directory, exist_ok=True)
        
        from document_creation import generate_resume_style_5
        docx_resume_path = generate_resume_style_5(json_resume_path, resume_directory, list_of_keywords=keywords_list if keywords_list else None)
        
        if not docx_resume_path or not os.path.exists(docx_resume_path):
            return jsonify({'error': 'Failed to create resume document'}), 500
        
        print(f"✅ Resume DOCX created at: {docx_resume_path}")
        
        # Return success with download path
        return jsonify({
            'success': True, 
            'resume_path': docx_resume_path,
            'message': 'Resume generated successfully!'
        }), 200
        
    except Exception as e:
        print(f"❌ Error generating resume from JSON: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/generate_resume_from_json_advanced', methods=['POST'])
def generate_resume_from_json_advanced():
    """Generate resume from JSON with support for .NET format and PDF conversion"""
    try:
        resume_json_str = request.form.get('resume_json', '')
        format_type = request.form.get('format', 'both')  # both, docx, pdf
        
        if not resume_json_str:
            return jsonify({'error': 'JSON resume data is required'}), 400
        
        # Validate JSON
        try:
            resume_json = json.loads(resume_json_str)
        except json.JSONDecodeError as e:
            return jsonify({'error': f'Invalid JSON: {str(e)}'}), 400
        
        # Load keywords from bold_words.json automatically
        keywords_list = []
        try:
            bold_words_path = os.path.join(os.getcwd(), 'bold_words.json')
            if os.path.exists(bold_words_path):
                with open(bold_words_path, 'r') as f:
                    bold_words_data = json.load(f)
                    if isinstance(bold_words_data, dict) and 'keywords' in bold_words_data:
                        keywords_list = bold_words_data['keywords']
                        print(f"✅ Loaded {len(keywords_list)} keywords from bold_words.json")
                    elif isinstance(bold_words_data, list):
                        keywords_list = bold_words_data
                        print(f"✅ Loaded {len(keywords_list)} keywords from bold_words.json (list format)")
        except Exception as e:
            print(f"⚠️ Could not load keywords from bold_words.json: {e}")
        
        # Determine which format to use based on JSON structure
        use_dotnet_format = 'personal_info' in resume_json or ('experience' in resume_json and resume_json['experience'] and 'project' in resume_json['experience'][0])
        
        resume_directory = "generated_resumes"
        os.makedirs(resume_directory, exist_ok=True)
        
        docx_path = None
        pdf_path = None
        
        # Generate DOCX
        if format_type in ['both', 'docx']:
            print("\n=== Converting JSON to DOCX ===")
            if use_dotnet_format:
                from document_creation import generate_resume_dotnet_format
                docx_path = generate_resume_dotnet_format(
                    resume_json_str, 
                    resume_directory, 
                    font_style="Times New Roman",
                    list_of_keywords=keywords_list if keywords_list else None
                )
            else:
                # Save JSON to temp file for style_5
                output_dir = "resumes"
                os.makedirs(output_dir, exist_ok=True)
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                json_filename = f"resume_from_json_{timestamp}.json"
                json_resume_path = os.path.join(output_dir, json_filename)
                
                with open(json_resume_path, 'w') as f:
                    json.dump(resume_json, f, indent=2)
                
                from document_creation import generate_resume_style_5
                docx_path = generate_resume_style_5(
                    json_resume_path, 
                    resume_directory, 
                    list_of_keywords=keywords_list if keywords_list else None
                )
            
            if not docx_path or not os.path.exists(docx_path):
                return jsonify({'error': 'Failed to create DOCX document'}), 500
            
            print(f"✅ Resume DOCX created at: {docx_path}")
        
        # Generate PDF if requested
        if format_type in ['both', 'pdf']:
            print("\n=== Converting DOCX to PDF ===")
            if not docx_path:
                return jsonify({'error': 'DOCX must be generated first to create PDF'}), 400
            
            try:
                # Try using docx2pdf (requires LibreOffice or Microsoft Word)
                try:
                    from docx2pdf import convert
                    pdf_path = docx_path.replace('.docx', '.pdf')
                    convert(docx_path, pdf_path)
                    print(f"✅ Resume PDF created at: {pdf_path}")
                except ImportError:
                    print("⚠️ docx2pdf not available, trying alternative method")
                    # Alternative: Use LibreOffice command line (if available)
                    import subprocess
                    pdf_path = docx_path.replace('.docx', '.pdf')
                    try:
                        subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(pdf_path), docx_path], 
                                     check=True, capture_output=True)
                        print(f"✅ Resume PDF created at: {pdf_path}")
                    except (subprocess.CalledProcessError, FileNotFoundError):
                        print("⚠️ PDF conversion not available. Please install LibreOffice or docx2pdf library.")
                        pdf_path = None
            except Exception as e:
                print(f"⚠️ PDF conversion failed: {e}")
                pdf_path = None
        
        result = {
            'success': True,
            'message': 'Resume generated successfully!'
        }
        
        if docx_path:
            result['docx_path'] = docx_path
        if pdf_path:
            result['pdf_path'] = pdf_path
        
        return jsonify(result), 200
        
    except Exception as e:
        print(f"❌ Error generating resume: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/download_resume', methods=['GET'])
def download_resume():
    """Download the generated resume"""
    try:
        resume_path = request.args.get('path', '')
        if not resume_path:
            return jsonify({'error': 'Resume file path is required'}), 400
        
        requested_path = os.path.realpath(resume_path)
        base_dir = os.path.realpath(os.getcwd())
        
        # Get desktop path for validation
        if os.name == 'nt':  # Windows
            desktop_path = os.path.realpath(os.path.join(os.path.expanduser('~'), 'Desktop'))
        else:  # macOS/Linux
            desktop_path = os.path.realpath(os.path.join(os.path.expanduser('~'), 'Desktop'))
        
        # Allow files from base directory or desktop
        if not (requested_path.startswith(base_dir) or requested_path.startswith(desktop_path)):
            return jsonify({'error': 'Invalid resume path'}), 400
        
        if not os.path.exists(requested_path):
            return jsonify({'error': 'Resume file not found'}), 404
        
        return send_file(requested_path, as_attachment=False, download_name=os.path.basename(requested_path))
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/preview_resume', methods=['GET'])
def preview_resume():
    """Stream resume file for in-browser preview (PDF/DOCX/DOC)."""
    try:
        resume_path = request.args.get('path', '')
        if not resume_path:
            return jsonify({'error': 'Resume file path is required'}), 400

        requested_path = os.path.realpath(resume_path)
        base_dir = os.path.realpath(os.getcwd())
        
        # Get desktop path for validation
        if os.name == 'nt':  # Windows
            desktop_path = os.path.realpath(os.path.join(os.path.expanduser('~'), 'Desktop'))
        else:  # macOS/Linux
            desktop_path = os.path.realpath(os.path.join(os.path.expanduser('~'), 'Desktop'))
        
        # Allow files from base directory or desktop
        if not (requested_path.startswith(base_dir) or requested_path.startswith(desktop_path)):
            return jsonify({'error': 'Invalid resume path'}), 400

        if not os.path.exists(requested_path):
            return jsonify({'error': 'Resume file not found'}), 404

        allowed_extensions = ('.pdf', '.docx', '.doc')
        if not requested_path.lower().endswith(allowed_extensions):
            return jsonify({'error': 'Unsupported file type for preview'}), 400

        mime_type = 'application/octet-stream'
        if requested_path.lower().endswith('.pdf'):
            mime_type = 'application/pdf'
        elif requested_path.lower().endswith('.docx'):
            mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif requested_path.lower().endswith('.doc'):
            mime_type = 'application/msword'

        return send_file(
            requested_path,
            mimetype=mime_type,
            as_attachment=False,
            download_name=os.path.basename(requested_path)
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/scrape_jobs', methods=['POST'])
def scrape_jobs():
    """Scrape job listings from website"""
    try:
        data = request.get_json()
        url = data.get('url', '')
        keywords = data.get('keywords', '')
        
        if not url:
            return jsonify({'error': 'URL is required'}), 400
        
        scraper = JobScraper()
        jobs = scraper.scrape_nvoids(url, keywords)
        
        return jsonify({'jobs': jobs, 'count': len(jobs)}), 200
        
    except Exception as e:
        print(f"❌ Error scraping jobs: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/available_senders', methods=['GET'])
def available_senders():
    """Return available sender emails (no passwords) from email.json"""
    try:
        email_json_path = os.path.join(os.getcwd(), 'email.json')
        if not os.path.exists(email_json_path):
            return jsonify({'senders': []}), 200

        with open(email_json_path, 'r') as f:
            data = json.load(f)

        senders = []
        # data is expected to be { email: [ {name, phone_number, smtp_username, smtp_password}, ... ] }
        for email, entries in data.items():
            if isinstance(entries, list) and len(entries) > 0 and isinstance(entries[0], dict):
                e = entries[0]
                senders.append({
                    'email': email,
                    'name': e.get('name', ''),
                    'phone': e.get('phone_number', ''),
                    'linkedin': e.get('linkedin_url', '') or ''
                })
            else:
                senders.append({'email': email, 'name': '', 'phone': '', 'linkedin': ''})

        return jsonify({'senders': senders}), 200
    except Exception as e:
        print('❌ Error reading email.json:', e)
        return jsonify({'senders': []}), 200

@app.route('/send_email', methods=['POST'])
def send_email():
    """Send email with resume attachment using EmailSender class"""
    temp_path = None
    try:
        # Get form data
        recruiter_name = request.form.get('recruiter_name', 'Hiring Manager')
        recruiter_email = request.form.get('recruiter_email')
        subject = request.form.get('subject', 'Application for Position')
        body = request.form.get('body', '')
        cc = request.form.get('cc', '').strip()  # Optional CC
        bcc = request.form.get('bcc', '').strip()  # Optional BCC
        resume_file = request.files.get('resume')
        resume_path = request.form.get('resume_path', '')
        
        if not recruiter_email:
            return jsonify({'success': False, 'error': 'Recruiter email is required'}), 400
        
        if not body:
            return jsonify({'success': False, 'error': 'Email body is required'}), 400
        
        # Handle resume attachment - check for server-side path or uploaded file
        if resume_path and os.path.exists(resume_path):
            # Use the server-side generated resume
            email_sender.resume_file = resume_path
            email_sender._original_filename = os.path.basename(resume_path)
            attach_resume = True
        elif resume_file:
            # Save uploaded resume temporarily
            original_filename = resume_file.filename
            print(f"📎 Original filename from upload: {original_filename}")
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            temp_filename = f"resume_{timestamp}_{original_filename}"
            temp_path = os.path.join(os.getcwd(), temp_filename)
            resume_file.save(temp_path)
            
            # Override filename in EmailSender to use original name (not the temp filename)
            email_sender.resume_file = temp_path
            email_sender._original_filename = original_filename
            print(f"✅ Preserving original filename for attachment: {original_filename}")
            attach_resume = True
        else:
            attach_resume = False
            email_sender.resume_file = ''
            if hasattr(email_sender, '_original_filename'):
                delattr(email_sender, '_original_filename')
        
        # Determine which sender to use. If frontend provided a sender_email, try to lookup its SMTP password
        sender_override = request.form.get('sender_email') or request.form.get('from_email')

        # Prepare metadata archive entry before attempting to send
        cc_list_normalised = _normalise_recipient_list(cc)
        bcc_list_normalised = _normalise_recipient_list(bcc)
        
        # Extract phone number from JD if available (from form data or body)
        raw_jd_text = request.form.get('raw_jd_text', '')  # JD text if passed
        phone_number = None
        
        print(f"\n=== Phone Extraction Debug ===")
        print(f"Raw JD text length: {len(raw_jd_text) if raw_jd_text else 0}")
        
        if raw_jd_text:
            # Try to extract phone number from JD - PRIORITY: Phone number only, not email
            phone_number = extract_phone_from_text(raw_jd_text)
            
            if phone_number:
                print(f"✅ Phone number extracted: {phone_number}")
            else:
                print(f"⚠️ No phone number found in JD text")
                # DO NOT save email in phone field - keep it null if no phone found
                phone_number = None
        
        # DO NOT use email as fallback - only save actual phone numbers
        if not phone_number:
            phone_number = None
            print(f"ℹ️ Phone field will be null (no phone number found)")

        email_archive_entry = {
            "to_email": recruiter_email,
            "name": recruiter_name,
            "subject": subject,
            "from_email": sender_override or email_sender.sender_email or '',
            "body": body,
            "timestamp": datetime.utcnow().isoformat(),
            "phone": phone_number if phone_number else None
        }

        if cc_list_normalised:
            email_archive_entry["cc"] = cc_list_normalised
        if bcc_list_normalised:
            email_archive_entry["bcc"] = bcc_list_normalised
        email_archive_entry["resume_attached"] = bool(attach_resume and (email_sender.resume_file or resume_path or resume_file))

        archive_email_metadata(email_archive_entry)

        if sender_override:
            # Load email.json from repo root and try to find matching entry
            try:
                print(f"\n=== Using sender email from selection: {sender_override} ===")
                email_json_path = os.path.join(os.getcwd(), 'email.json')
                if os.path.exists(email_json_path):
                    with open(email_json_path, 'r') as ef:
                        emails_data = json.load(ef)
                    print(f"✅ Loaded email.json with {len(emails_data)} email entries")
                else:
                    emails_data = {}
                    print(f"⚠️ email.json not found at {email_json_path}")

                # emails_data format: { "email@x.com": [ { name, phone_number, smtp_username, smtp_password } ] }
                smtp_pw = None
                sender_name = None
                if sender_override in emails_data and isinstance(emails_data[sender_override], list) and len(emails_data[sender_override])>0:
                    first = emails_data[sender_override][0]
                    smtp_pw = first.get('smtp_password') or first.get('app_password')
                    sender_name = first.get('name')
                    print(f"✅ Found SMTP credentials for {sender_override} (Name: {sender_name})")
                else:
                    print(f"⚠️ Email {sender_override} not found in email.json")

                # Fallback to environment if password not found
                if not smtp_pw:
                    smtp_pw = os.getenv('SMTP_PASSWORD', '')
                    print(f"⚠️ Using fallback SMTP password from environment variables")

                if not smtp_pw:
                    print(f"❌ No SMTP password found for {sender_override}")
                    return jsonify({'success': False, 'error': f'No SMTP password found for selected email: {sender_override}'}), 400

                # Create a temporary EmailSender for this send so we don't mutate global
                temp_sender = EmailSender(sender_email=sender_override, app_password=smtp_pw, resume_file=email_sender.resume_file)
                # Preserve the original filename so attachment uses original name, not temp filename
                if hasattr(email_sender, '_original_filename'):
                    temp_sender._original_filename = email_sender._original_filename
                print(f"📧 Sending email using {sender_override}")

                success = temp_sender.send_email(
                    recruiter_email=recruiter_email,
                    recruiter_name=recruiter_name,
                    subject=subject,
                    body=body,
                    attach_resume=attach_resume,
                    cc=cc if cc else None,
                    bcc=bcc if bcc else None
                )
            except Exception as e:
                print(f'❌ Error preparing sender override: {e}')
                import traceback
                traceback.print_exc()
                success = False
        else:
            # Use default configured email_sender instance
            success = email_sender.send_email(
                recruiter_email=recruiter_email,
                recruiter_name=recruiter_name,
                subject=subject,
                body=body,
                attach_resume=attach_resume,
                cc=cc if cc else None,
                bcc=bcc if bcc else None
            )
        
        if success:
            print(f"✅ Email sent successfully to {recruiter_email}")
            return jsonify({'success': True, 'message': 'Email sent successfully'}), 200
        else:
            return jsonify({'success': False, 'error': 'Failed to send email'}), 500
        
    except Exception as e:
        print(f"❌ Error sending email: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500
    
    finally:
        # Clean up temp file if created
        if temp_path and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except:
                pass
        if hasattr(email_sender, '_original_filename'):
            delattr(email_sender, '_original_filename')

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5002)

