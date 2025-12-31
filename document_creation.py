import json
import os
import re
from datetime import datetime
from typing import List, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import qn
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def load_bold_keywords(bold_words_file: str = "bold_words.json") -> List[str]:
    """Load keywords from bold_words.json file"""
    try:
        if os.path.exists(bold_words_file):
            with open(bold_words_file, 'r') as f:
                data = json.load(f)
                # Handle both list and dict formats
                if isinstance(data, list):
                    return data
                elif isinstance(data, dict) and 'keywords' in data:
                    return data['keywords']
                elif isinstance(data, dict):
                    # Try to extract any values that look like keywords
                    keywords = []
                    for key, value in data.items():
                        if isinstance(value, str):
                            keywords.append(value)
                        elif isinstance(value, list):
                            keywords.extend(value)
                    return keywords
        return []
    except Exception as e:
        print(f"Error loading bold keywords: {e}")
        return []

def generate_resume_style_1(json_file_path: str, resume_directory: str, font_style: str = "Calibri", list_of_keywords: Optional[List[str]] = None) -> str:
    """
    Generate resume from JSON with style_2 format: Personal Information, Professional Summary, 
    Professional Experience, Technical Skills, Education, Certifications
    
    Args:
        json_file_path: Path to JSON file with resume data
        resume_directory: Directory to save the generated resume
        font_style: Font name to use for the resume (default: "Calibri")
        list_of_keywords: List of keywords to highlight in the resume (optional)
    
    Returns:
        Path to generated resume file
    """
    # Load JSON data
    with open(json_file_path, 'r') as file:
        data = json.load(file)
    name = data.get('name', 'resume').replace(" ", "_")
    # Add timestamp to prevent file overriding
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # Create today's date folder (format: YYYY-MM-DD)
    today_folder = datetime.now().strftime("%Y-%m-%d")
    today_resume_directory = os.path.join(resume_directory, today_folder)
    
    # Ensure the today's folder exists
    os.makedirs(today_resume_directory, exist_ok=True)
    
    output_path = os.path.join(today_resume_directory, f"{name}_CV.docx")
        
    doc = Document()
    # Set tight margins (0.5 inches)
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    # Add page border (0.2 inches from page edges)
    sectPr = section._sectPr
    pgBorders = parse_xml(r'<w:pgBorders {0} w:offsetFrom="page"><w:top w:val="single" w:sz="6" w:space="14.4" w:color="000000"/><w:left w:val="single" w:sz="6" w:space="14.4" w:color="000000"/><w:bottom w:val="single" w:sz="6" w:space="14.4" w:color="000000"/><w:right w:val="single" w:sz="6" w:space="14.4" w:color="000000"/></w:pgBorders>'.format(nsdecls('w')))
    sectPr.append(pgBorders)
    
    # Configure base style with provided font
    style = doc.styles['Normal']
    font = style.font  # type: ignore
    font.name = font_style
    font.size = Pt(11)
    style.paragraph_format.space_before = Pt(0)  # type: ignore
    style.paragraph_format.space_after = Pt(0)  # type: ignore
    style.paragraph_format.line_spacing = 1.0  # type: ignore
    # Load bold keywords from bold_words.json
    bold_keywords = load_bold_keywords()
    
    # Extract all skills for highlighting (only if list_of_keywords is provided)
    all_skills = []
    if list_of_keywords:
        # Add technical skills from the resume
        if 'technical_skills' in data and data['technical_skills']:
            for category, skills in data['technical_skills'].items():
                all_skills.extend(skills)
        
        # Add provided keywords to skills list
        all_skills.extend(list_of_keywords)
    
    # Always add bold keywords to the skills list for bolding
    all_skills.extend(bold_keywords)
    # Helper function to add hyperlinks with blue color
    def add_hyperlink(paragraph, text, url):
        """Add a hyperlink to a paragraph with blue color"""
        # Create hyperlink element
        hyperlink = parse_xml(r'<w:hyperlink {0} w:history="1"><w:r><w:rPr><w:rStyle w:val="Hyperlink"/></w:rPr><w:t>{1}</w:t></w:r></w:hyperlink>'.format(nsdecls('w'), text))
        hyperlink.set(qn('r:id'), paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True))
        paragraph._p.append(hyperlink)
        
        # Set the hyperlink color to blue
        for run in hyperlink.iter():
            if run.tag.endswith('}r'):
                rPr = run.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
                if rPr is not None:
                    # Add color element for blue
                    color_elem = parse_xml(r'<w:color {0} w:val="0000FF"/>'.format(nsdecls('w')))
                    rPr.append(color_elem)
                    # Add underline element
                    underline_elem = parse_xml(r'<w:u {0} w:val="single"/>'.format(nsdecls('w')))
                    rPr.append(underline_elem)
        
        return hyperlink
    # Helper functions
    def add_centered_paragraph(text, bold=False, size=11, color=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        return p
    def add_justified_paragraph(text=None, bold=False, color=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if text:
            run = p.add_run(text)
            run.bold = bold
            if color:
                run.font.color.rgb = RGBColor(*color)
        return p
    def add_section_heading(text):
        # Add bold section heading with border line (no gap)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text.upper())
        run.bold = True
        run.font.name = font_style
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add border line directly under text (no gap)
        p_border = parse_xml(r'<w:pBdr {0}><w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/></w:pBdr>'.format(nsdecls('w')))
        p._p.get_or_add_pPr().append(p_border)
        
        p.paragraph_format.space_after = Pt(4)
    def add_bullet_points(items):
        for item in items:
            p = doc.add_paragraph(style='List Bullet')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            if all_skills:
                highlight_skills_in_text(item, p, all_skills)
            else:
                run = p.add_run(item)
                run.font.size = Pt(11)
    def highlight_skills_in_text(text, paragraph, skills_list):
        """Highlight skills in text with bold formatting using exact keyword matching"""
        if not skills_list:
            run = paragraph.add_run(text)
            run.font.size = Pt(11)
            return
        
        # Create exact match patterns with word boundaries
        exact_patterns = []
        for skill in skills_list:
            # Escape special regex characters
            escaped_skill = re.escape(skill.strip())
            # Create pattern with word boundaries for exact matching
            # This ensures the skill is matched as a complete word/phrase
            pattern = r'\b' + escaped_skill + r'\b'
            exact_patterns.append((pattern, skill.strip()))
        
        # Find all matches with their positions
        matches = []
        for pattern, original_skill in exact_patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                matches.append((match.start(), match.end(), match.group(), original_skill))
        
        # Sort matches by position
        matches.sort(key=lambda x: x[0])
        
        # Remove overlapping matches (keep the first one)
        filtered_matches = []
        for match in matches:
            if not any(existing[0] <= match[0] < existing[1] or existing[0] < match[1] <= existing[1] 
                      for existing in filtered_matches):
                filtered_matches.append(match)
        
        if not filtered_matches:
            run = paragraph.add_run(text)
            run.font.size = Pt(11)
            return
        
        paragraph.clear()
        last_end = 0
        
        for start, end, matched_text, original_skill in filtered_matches:
            # Add text before the match
            if start > last_end:
                run = paragraph.add_run(text[last_end:start])
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Add the highlighted skill
            skill_run = paragraph.add_run(matched_text)
            skill_run.bold = True
            skill_run.font.color.rgb = RGBColor(0, 0, 0)
            skill_run.font.size = Pt(11)
            
            last_end = end
        
        # Add remaining text after the last match
        if last_end < len(text):
            run = paragraph.add_run(text[last_end:])
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)
    
    # --- Personal Information (Name, Title, Contact) ---
    add_centered_paragraph(data['name'], bold=True, size=16)
    add_centered_paragraph(data['title'], size=12)
    contact = data['contact']
    # Create contact paragraph with hyperlinks
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    first_item = True
    
    # Email with mailto: hyperlink
    if 'email' in contact:
        if not first_item:
            contact_p.add_run(" | ")
        add_hyperlink(contact_p, contact['email'], f"mailto:{contact['email']}")
        first_item = False
    # Phone with tel: hyperlink
    if 'phone' in contact:
        if not first_item:
            contact_p.add_run(" | ")
        add_hyperlink(contact_p, contact['phone'], f"tel:{contact['phone']}")
        first_item = False
    # LinkedIn with URL hyperlink
    if 'linkedin' in contact and contact['linkedin'] and contact['linkedin'].strip():
        if not first_item:
            contact_p.add_run(" | ")
        add_hyperlink(contact_p, "LinkedIn", contact['linkedin'])
        first_item = False
    # Portfolio with URL hyperlink
    if contact.get('portfolio'):
        if not first_item:
            contact_p.add_run(" | ")
        add_hyperlink(contact_p, "Portfolio", contact['portfolio'])
    
    # --- Professional Summary ---
    if 'professional_summary' in data and data['professional_summary']:
        add_section_heading("Professional Summary")
        add_bullet_points(data['professional_summary'])
    
    # --- Professional Experience ---
    if 'experience' in data and data['experience']:
        add_section_heading("Professional Experience")
        for i, job in enumerate(data['experience']):
            # Create paragraph for Job Title with Duration right-aligned
            role_para = doc.add_paragraph()
            role_para.paragraph_format.space_after = Pt(0)
            
            # Add job title (left-aligned, bold)
            title_run = role_para.add_run(job['role'])
            title_run.bold = True
            title_run.font.color.rgb = RGBColor(0, 0, 0)
            title_run.font.size = Pt(11)
            
            # Add tab stops for right alignment of duration
            tab_stops = role_para.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
            
            # Add tab and duration (right-aligned, bold)
            duration_run = role_para.add_run("\t" + job['duration'])
            duration_run.bold = True
            duration_run.font.color.rgb = RGBColor(0, 0, 0)
            duration_run.font.size = Pt(11)
            
            # Client and location information (next line, left-aligned)
            client = job.get('client', '').strip()
            location = job.get('location', '').strip() if 'location' in job else ''
            if client and location:
                client_info = f"{client} – {location}"
            elif client:
                client_info = client
            elif location:
                client_info = location
            else:
                client_info = ''
                
            if client_info:
                client_para = doc.add_paragraph()
                client_para.paragraph_format.space_before = Pt(0)
                client_para.paragraph_format.space_after = Pt(2)
                client_run = client_para.add_run(client_info)
                client_run.bold = True
                client_run.font.color.rgb = RGBColor(0, 0, 0)
                client_run.font.size = Pt(11)
                
            if 'responsibilities' in job and job['responsibilities']:
                if job.get('client') == "Cipla Pharmaceuticals":
                    p = add_justified_paragraph("Responsibilities: To modernize Cipla's healthcare data infrastructure", bold=True)
                else:
                    p = add_justified_paragraph("Responsibilities:", bold=True)
                add_bullet_points(job['responsibilities'])
                
            if 'environment' in job and job['environment'] and len(job['environment']) > 0:
                p = add_justified_paragraph("Environment:", bold=True)
                env_text = ", ".join(job['environment'])
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Environment section: No bold skills highlighting, just plain text
                run = p.add_run(env_text)
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Add minimal spacing between clients
            if i < len(data['experience']) - 1:
                spacing_para = doc.add_paragraph()
                spacing_para.paragraph_format.space_after = Pt(2)
    
    # --- Technical Skills ---
    if 'technical_skills' in data and data['technical_skills']:
        add_section_heading("Technical Skills")
        first_category = True
        for category, skills in data['technical_skills'].items():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Remove extra spacing for first category
            if first_category:
                p.paragraph_format.space_before = Pt(0)
                first_category = False
            run = p.add_run(f"{category}: ")
            run.font.size = Pt(11)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            
            skills_text = ", ".join(skills)
            skills_run = p.add_run(skills_text)
            skills_run.font.size = Pt(11)
            skills_run.bold = False
    
    # --- Education ---
    if 'education' in data:
        add_section_heading("Education")
        educations = data['education'] if isinstance(data['education'], list) else [data['education']]
        for edu in educations:
            p = doc.add_paragraph()
            if 'degree' in edu:
                degree_run = p.add_run(f"{edu['degree']} ")
                degree_run.bold = True
                degree_run.font.color.rgb = RGBColor(0, 0, 0)
                remaining_parts = []
                if 'field' in edu:
                    remaining_parts.append(f"in {edu['field']}")
                if 'concentration' in edu and edu['concentration']:
                    remaining_parts.append(f"({edu['concentration']})")
                if 'institution' in edu:
                    remaining_parts.append(f"| {edu['institution']}")
                if 'year' in edu and edu['year'] and str(edu['year']).strip().upper() != 'YYYY':
                    remaining_parts.append(f"({edu['year']})")
                remaining_text = " ".join(remaining_parts)
                p.add_run(remaining_text)
    
    # --- Certifications ---
    if 'certifications' in data and data['certifications'] and len(data['certifications']) > 0:
        add_section_heading("Certifications")
        # Don't bold keywords in certifications
        for item in data['certifications']:
            p = doc.add_paragraph(style='List Bullet')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(item)
            run.font.size = Pt(11)
    
    doc.save(output_path)
    return output_path


def generate_resume_style_2(json_file_path: str, resume_directory: str, font_style: str = "Calibri", list_of_keywords: Optional[List[str]] = None) -> str:
    """
    Generate resume from JSON with style_2 format: Personal Information, Professional Summary, 
    Professional Experience, Technical Skills, Education, Certifications
    WITH CONDITIONAL STYLING: Border lines for Professional Summary & Technical Skills, 
    Light blue shading for other sections
    
    Args:
        json_file_path: Path to JSON file with resume data
        resume_directory: Directory to save the generated resume
        font_style: Font name to use for the resume (default: "Calibri")
        list_of_keywords: List of keywords to highlight in the resume (optional)
    
    Returns:
        Path to generated resume file
    """
    # Load JSON data
    with open(json_file_path, 'r') as file:
        data = json.load(file)
    name = data.get('name', 'resume').replace(" ", "_")
    # Add timestamp to prevent file overriding
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # Create today's date folder (format: YYYY-MM-DD)
    today_folder = datetime.now().strftime("%Y-%m-%d")
    today_resume_directory = os.path.join(resume_directory, today_folder)
    
    # Ensure the today's folder exists
    os.makedirs(today_resume_directory, exist_ok=True)
    
    output_path = os.path.join(today_resume_directory, f"{name}_style_2_conditional_{timestamp}.docx")
        
    doc = Document()
    # Set tight margins (0.5 inches)
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    # Add page border (0.2 inches from page edges)
    sectPr = section._sectPr
    pgBorders = parse_xml(r'<w:pgBorders {0} w:offsetFrom="page"><w:top w:val="single" w:sz="6" w:space="14.4" w:color="000000"/><w:left w:val="single" w:sz="6" w:space="14.4" w:color="000000"/><w:bottom w:val="single" w:sz="6" w:space="14.4" w:color="000000"/><w:right w:val="single" w:sz="6" w:space="14.4" w:color="000000"/></w:pgBorders>'.format(nsdecls('w')))
    sectPr.append(pgBorders)
    
    # Configure base style with provided font
    style = doc.styles['Normal']
    font = style.font  # type: ignore
    font.name = font_style
    font.size = Pt(11)
    style.paragraph_format.space_before = Pt(0)  # type: ignore
    style.paragraph_format.space_after = Pt(0)  # type: ignore
    style.paragraph_format.line_spacing = 1.0  # type: ignore
    # Load bold keywords from bold_words.json
    bold_keywords = load_bold_keywords()
    
    # Extract all skills for highlighting (only if list_of_keywords is provided)
    all_skills = []
    if list_of_keywords:
        # Add technical skills from the resume
        if 'technical_skills' in data and data['technical_skills']:
            for category, skills in data['technical_skills'].items():
                all_skills.extend(skills)
        
        # Add provided keywords to skills list
        all_skills.extend(list_of_keywords)
    
    # Always add bold keywords to the skills list for bolding
    all_skills.extend(bold_keywords)
    # Helper function to add hyperlinks with blue color
    def add_hyperlink(paragraph, text, url):
        """Add a hyperlink to a paragraph with blue color"""
        # Create hyperlink element
        hyperlink = parse_xml(r'<w:hyperlink {0} w:history="1"><w:r><w:rPr><w:rStyle w:val="Hyperlink"/></w:rPr><w:t>{1}</w:t></w:r></w:hyperlink>'.format(nsdecls('w'), text))
        hyperlink.set(qn('r:id'), paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True))
        paragraph._p.append(hyperlink)
        
        # Set the hyperlink color to blue
        for run in hyperlink.iter():
            if run.tag.endswith('}r'):
                rPr = run.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
                if rPr is not None:
                    # Add color element for blue
                    color_elem = parse_xml(r'<w:color {0} w:val="0000FF"/>'.format(nsdecls('w')))
                    rPr.append(color_elem)
                    # Add underline element
                    underline_elem = parse_xml(r'<w:u {0} w:val="single"/>'.format(nsdecls('w')))
                    rPr.append(underline_elem)
        
        return hyperlink
    # Helper functions
    def add_centered_paragraph(text, bold=False, size=11, color=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        return p
    def add_justified_paragraph(text=None, bold=False, color=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if text:
            run = p.add_run(text)
            run.bold = bold
            if color:
                run.font.color.rgb = RGBColor(*color)
        return p
    def add_section_heading_conditional(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text.upper())
        run.bold = True
        run.font.name = font_style
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Apply light blue shading to ALL sections (remove separation line)
        shading = parse_xml(r'<w:shd {0} w:val="clear" w:color="auto" w:fill="ADD8E6"/>'.format(nsdecls('w')))
        p._p.get_or_add_pPr().append(shading)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
    def add_bullet_points(items):
        for item in items:
            p = doc.add_paragraph(style='List Bullet')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            if all_skills:
                highlight_skills_in_text(item, p, all_skills)
            else:
                run = p.add_run(item)
                run.font.size = Pt(11)
    def highlight_skills_in_text(text, paragraph, skills_list):
        """Highlight skills in text with bold formatting using exact keyword matching"""
        if not skills_list:
            run = paragraph.add_run(text)
            run.font.size = Pt(11)
            return
        
        # Create exact match patterns with word boundaries
        exact_patterns = []
        for skill in skills_list:
            # Escape special regex characters
            escaped_skill = re.escape(skill.strip())
            # Create pattern with word boundaries for exact matching
            # This ensures the skill is matched as a complete word/phrase
            pattern = r'\b' + escaped_skill + r'\b'
            exact_patterns.append((pattern, skill.strip()))
        
        # Find all matches with their positions
        matches = []
        for pattern, original_skill in exact_patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                matches.append((match.start(), match.end(), match.group(), original_skill))
        
        # Sort matches by position
        matches.sort(key=lambda x: x[0])
        
        # Remove overlapping matches (keep the first one)
        filtered_matches = []
        for match in matches:
            if not any(existing[0] <= match[0] < existing[1] or existing[0] < match[1] <= existing[1] 
                      for existing in filtered_matches):
                filtered_matches.append(match)
        
        if not filtered_matches:
            run = paragraph.add_run(text)
            run.font.size = Pt(11)
            return
        
        paragraph.clear()
        last_end = 0
        
        for start, end, matched_text, original_skill in filtered_matches:
            # Add text before the match
            if start > last_end:
                run = paragraph.add_run(text[last_end:start])
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Add the highlighted skill
            skill_run = paragraph.add_run(matched_text)
            skill_run.bold = True
            skill_run.font.color.rgb = RGBColor(0, 0, 0)
            skill_run.font.size = Pt(11)
            
            last_end = end
        
        # Add remaining text after the last match
        if last_end < len(text):
            run = paragraph.add_run(text[last_end:])
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)
    
    # --- Personal Information (Name, Title, Contact) ---
    add_centered_paragraph(data['name'], bold=True, size=16)
    add_centered_paragraph(data['title'], size=12)
    contact = data['contact']
    # Create contact paragraph with hyperlinks
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    first_item = True
    
    # Email with mailto: hyperlink
    if 'email' in contact:
        if not first_item:
            contact_p.add_run(" | ")
        add_hyperlink(contact_p, contact['email'], f"mailto:{contact['email']}")
        first_item = False
    # Phone with tel: hyperlink
    if 'phone' in contact:
        if not first_item:
            contact_p.add_run(" | ")
        add_hyperlink(contact_p, contact['phone'], f"tel:{contact['phone']}")
        first_item = False
    # LinkedIn with URL hyperlink
    if 'linkedin' in contact and contact['linkedin'] and contact['linkedin'].strip():
        if not first_item:
            contact_p.add_run(" | ")
        add_hyperlink(contact_p, "LinkedIn", contact['linkedin'])
        first_item = False
    # Portfolio with URL hyperlink
    if contact.get('portfolio'):
        if not first_item:
            contact_p.add_run(" | ")
        add_hyperlink(contact_p, "Portfolio", contact['portfolio'])
    
    # --- Professional Summary ---
    if 'professional_summary' in data and data['professional_summary']:
        add_section_heading_conditional("Professional Summary")
        add_bullet_points(data['professional_summary'])
    
    # --- Professional Experience ---
    if 'experience' in data and data['experience']:
        add_section_heading_conditional("Professional Experience")
        for i, job in enumerate(data['experience']):
            # Create paragraph for Job Title with Duration right-aligned
            role_para = doc.add_paragraph()
            role_para.paragraph_format.space_after = Pt(0)
            
            # Add job title (left-aligned, bold)
            title_run = role_para.add_run(job['role'])
            title_run.bold = True
            title_run.font.color.rgb = RGBColor(0, 0, 0)
            title_run.font.size = Pt(11)
            
            # Add tab stops for right alignment of duration
            tab_stops = role_para.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
            
            # Add tab and duration (right-aligned, bold)
            duration_run = role_para.add_run("\t" + job['duration'])
            duration_run.bold = True
            duration_run.font.color.rgb = RGBColor(0, 0, 0)
            duration_run.font.size = Pt(11)
            
            # Client and location information (next line, left-aligned)
            client = job.get('client', '').strip()
            location = job.get('location', '').strip() if 'location' in job else ''
            if client and location:
                client_info = f"{client} – {location}"
            elif client:
                client_info = client
            elif location:
                client_info = location
            else:
                client_info = ''
                
            if client_info:
                client_para = doc.add_paragraph()
                client_para.paragraph_format.space_before = Pt(0)
                client_para.paragraph_format.space_after = Pt(2)
                client_run = client_para.add_run(client_info)
                client_run.bold = True
                client_run.font.color.rgb = RGBColor(0, 0, 0)
                client_run.font.size = Pt(11)
                
            if 'responsibilities' in job and job['responsibilities']:
                if job.get('client') == "Cipla Pharmaceuticals":
                    p = add_justified_paragraph("Responsibilities: To modernize Cipla's healthcare data infrastructure", bold=True)
                else:
                    p = add_justified_paragraph("Responsibilities:", bold=True)
                add_bullet_points(job['responsibilities'])
                
            if 'environment' in job and job['environment'] and len(job['environment']) > 0:
                p = add_justified_paragraph("Environment:", bold=True)
                env_text = ", ".join(job['environment'])
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Environment section: No bold skills highlighting, just plain text
                run = p.add_run(env_text)
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Add minimal spacing between clients
            if i < len(data['experience']) - 1:
                spacing_para = doc.add_paragraph()
                spacing_para.paragraph_format.space_after = Pt(2)
    
    # --- Technical Skills ---
    if 'technical_skills' in data and data['technical_skills']:
        add_section_heading_conditional("Technical Skills")
        first_category = True
        for category, skills in data['technical_skills'].items():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Remove extra spacing for first category
            if first_category:
                p.paragraph_format.space_before = Pt(0)
                first_category = False
            run = p.add_run(f"{category}: ")
            run.font.size = Pt(11)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            
            skills_text = ", ".join(skills)
            skills_run = p.add_run(skills_text)
            skills_run.font.size = Pt(11)
            skills_run.bold = False
    
    # --- Education ---
    if 'education' in data:
        add_section_heading_conditional("Education")
        educations = data['education'] if isinstance(data['education'], list) else [data['education']]
        for edu in educations:
            p = doc.add_paragraph()
            if 'degree' in edu:
                degree_run = p.add_run(f"{edu['degree']} ")
                degree_run.bold = True
                degree_run.font.color.rgb = RGBColor(0, 0, 0)
                remaining_parts = []
                if 'field' in edu:
                    remaining_parts.append(f"in {edu['field']}")
                if 'concentration' in edu and edu['concentration']:
                    remaining_parts.append(f"({edu['concentration']})")
                if 'institution' in edu:
                    remaining_parts.append(f"| {edu['institution']}")
                if 'year' in edu and edu['year'] and str(edu['year']).strip().upper() != 'YYYY':
                    remaining_parts.append(f"({edu['year']})")
                remaining_text = " ".join(remaining_parts)
                p.add_run(remaining_text)
    
    # --- Certifications ---
    if 'certifications' in data and data['certifications'] and len(data['certifications']) > 0:
        add_section_heading_conditional("Certifications")
        # Don't bold keywords in certifications
        for item in data['certifications']:
            p = doc.add_paragraph(style='List Bullet')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(item)
            run.font.size = Pt(11)
    
    doc.save(output_path)
    return output_path


def generate_resume_style_3(json_file_path: str, resume_directory: str, font_style: str = "Calibri", list_of_keywords: Optional[List[str]] = None) -> str:
    """
    Generate resume from JSON with style_3 format: Personal Information, Professional Summary, 
    Professional Experience, Technical Skills, Education, Certifications
    WITH ENHANCED ROLE AND CLIENT METADATA DISPLAY
    
    Args:
        json_file_path: Path to JSON file with resume data
        resume_directory: Directory to save the generated resume
        font_style: Font name to use for the resume (default: "Calibri")
        list_of_keywords: List of keywords to highlight in the resume (optional)
    
    Returns:
        Path to generated resume file
    """
    # Load JSON data
    with open(json_file_path, 'r') as file:
        data = json.load(file)
    name = data.get('name', 'resume').replace(" ", "_")
    # Add timestamp to prevent file overriding
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # Create today's date folder (format: YYYY-MM-DD)
    today_folder = datetime.now().strftime("%Y-%m-%d")
    today_resume_directory = os.path.join(resume_directory, today_folder)
    
    # Ensure the today's folder exists
    os.makedirs(today_resume_directory, exist_ok=True)
    
    output_path = os.path.join(today_resume_directory, f"{name}_style_3_enhanced_{timestamp}.docx")
        
    doc = Document()
    # Set tight margins (0.5 inches)
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    # Add page border (0.2 inches from page edges)
    sectPr = section._sectPr
    pgBorders = parse_xml(r'<w:pgBorders {0} w:offsetFrom="page"><w:top w:val="single" w:sz="6" w:space="14.4" w:color="000000"/><w:left w:val="single" w:sz="6" w:space="14.4" w:color="000000"/><w:bottom w:val="single" w:sz="6" w:space="14.4" w:color="000000"/><w:right w:val="single" w:sz="6" w:space="14.4" w:color="000000"/></w:pgBorders>'.format(nsdecls('w')))
    sectPr.append(pgBorders)
    
    # Configure base style with provided font
    style = doc.styles['Normal']
    font = style.font  # type: ignore
    font.name = font_style
    font.size = Pt(11)
    style.paragraph_format.space_before = Pt(0)  # type: ignore
    style.paragraph_format.space_after = Pt(0)  # type: ignore
    style.paragraph_format.line_spacing = 1.0  # type: ignore
    # Load bold keywords from bold_words.json
    bold_keywords = load_bold_keywords()
    
    # Extract all skills for highlighting (only if list_of_keywords is provided)
    all_skills = []
    if list_of_keywords:
        # Add technical skills from the resume
        if 'technical_skills' in data and data['technical_skills']:
            for category, skills in data['technical_skills'].items():
                all_skills.extend(skills)
        
        # Add provided keywords to skills list
        all_skills.extend(list_of_keywords)
    
    # Always add bold keywords to the skills list for bolding
    all_skills.extend(bold_keywords)
    # Helper function to add hyperlinks with blue color
    def add_hyperlink(paragraph, text, url):
        """Add a hyperlink to a paragraph with blue color"""
        # Create hyperlink element
        hyperlink = parse_xml(r'<w:hyperlink {0} w:history="1"><w:r><w:rPr><w:rStyle w:val="Hyperlink"/></w:rPr><w:t>{1}</w:t></w:r></w:hyperlink>'.format(nsdecls('w'), text))
        hyperlink.set(qn('r:id'), paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True))
        paragraph._p.append(hyperlink)
        
        # Set the hyperlink color to blue
        for run in hyperlink.iter():
            if run.tag.endswith('}r'):
                rPr = run.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
                if rPr is not None:
                    # Add color element for blue
                    color_elem = parse_xml(r'<w:color {0} w:val="0000FF"/>'.format(nsdecls('w')))
                    rPr.append(color_elem)
                    # Add underline element
                    underline_elem = parse_xml(r'<w:u {0} w:val="single"/>'.format(nsdecls('w')))
                    rPr.append(underline_elem)
        
        return hyperlink
    # Helper functions
    def add_centered_paragraph(text, bold=False, size=11, color=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        return p
    def add_justified_paragraph(text=None, bold=False, color=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if text:
            run = p.add_run(text)
            run.bold = bold
            if color:
                run.font.color.rgb = RGBColor(*color)
        return p
    def add_section_heading_conditional(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text.upper())
        run.bold = True
        run.font.name = font_style
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Apply light blue shading to ALL sections (remove separation line)
        shading = parse_xml(r'<w:shd {0} w:val="clear" w:color="auto" w:fill="ADD8E6"/>'.format(nsdecls('w')))
        p._p.get_or_add_pPr().append(shading)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
    def add_bullet_points(items):
        for item in items:
            p = doc.add_paragraph(style='List Bullet')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            if all_skills:
                highlight_skills_in_text(item, p, all_skills)
            else:
                run = p.add_run(item)
                run.font.size = Pt(11)
    def highlight_skills_in_text(text, paragraph, skills_list):
        """Highlight skills in text with bold formatting using exact keyword matching"""
        if not skills_list:
            run = paragraph.add_run(text)
            run.font.size = Pt(11)
            return
        
        # Create exact match patterns with word boundaries
        exact_patterns = []
        for skill in skills_list:
            # Escape special regex characters
            escaped_skill = re.escape(skill.strip())
            # Create pattern with word boundaries for exact matching
            # This ensures the skill is matched as a complete word/phrase
            pattern = r'\b' + escaped_skill + r'\b'
            exact_patterns.append((pattern, skill.strip()))
        
        # Find all matches with their positions
        matches = []
        for pattern, original_skill in exact_patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                matches.append((match.start(), match.end(), match.group(), original_skill))
        
        # Sort matches by position
        matches.sort(key=lambda x: x[0])
        
        # Remove overlapping matches (keep the first one)
        filtered_matches = []
        for match in matches:
            if not any(existing[0] <= match[0] < existing[1] or existing[0] < match[1] <= existing[1] 
                      for existing in filtered_matches):
                filtered_matches.append(match)
        
        if not filtered_matches:
            run = paragraph.add_run(text)
            run.font.size = Pt(11)
            return
        
        paragraph.clear()
        last_end = 0
        
        for start, end, matched_text, original_skill in filtered_matches:
            # Add text before the match
            if start > last_end:
                run = paragraph.add_run(text[last_end:start])
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Add the highlighted skill
            skill_run = paragraph.add_run(matched_text)
            skill_run.bold = True
            skill_run.font.color.rgb = RGBColor(0, 0, 0)
            skill_run.font.size = Pt(11)
            
            last_end = end
        
        # Add remaining text after the last match
        if last_end < len(text):
            run = paragraph.add_run(text[last_end:])
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)
    
    # --- Personal Information (Name, Title, Contact) ---
    add_centered_paragraph(data['name'], bold=True, size=16)
    add_centered_paragraph(data['title'], size=12)
    contact = data['contact']
    # Create contact paragraph with hyperlinks
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    first_item = True
    
    # Email with mailto: hyperlink
    if 'email' in contact:
        if not first_item:
            contact_p.add_run(" | ")
        add_hyperlink(contact_p, contact['email'], f"mailto:{contact['email']}")
        first_item = False
    # Phone with tel: hyperlink
    if 'phone' in contact:
        if not first_item:
            contact_p.add_run(" | ")
        add_hyperlink(contact_p, contact['phone'], f"tel:{contact['phone']}")
        first_item = False
    # LinkedIn with URL hyperlink
    if 'linkedin' in contact and contact['linkedin'] and contact['linkedin'].strip():
        if not first_item:
            contact_p.add_run(" | ")
        add_hyperlink(contact_p, "LinkedIn", contact['linkedin'])
        first_item = False
    # Portfolio with URL hyperlink
    if contact.get('portfolio'):
        if not first_item:
            contact_p.add_run(" | ")
        add_hyperlink(contact_p, "Portfolio", contact['portfolio'])
    
    # --- Professional Summary ---
    if 'professional_summary' in data and data['professional_summary']:
        add_section_heading_conditional("Professional Summary")
        add_bullet_points(data['professional_summary'])
    
    # --- Professional Experience with Enhanced Role and Client Metadata ---
    if 'experience' in data and data['experience']:
        add_section_heading_conditional("Professional Experience")
        for i, job in enumerate(data['experience']):
            # Job title, company, and duration in single line format
            job_para = doc.add_paragraph()
            job_para.paragraph_format.space_after = Pt(2)
            
            # Format: "Role || Client || Duration"
            role = job.get('role', 'N/A')
            company = job.get('client', job.get('company', 'N/A'))
            duration = job.get('duration', 'N/A')
            
            job_run = job_para.add_run(f"{role} || {company} || {duration}")
            job_run.bold = True
            job_run.font.color.rgb = RGBColor(0, 0, 0)
            job_run.font.size = Pt(11)
                
            if 'responsibilities' in job and job['responsibilities']:
                if job.get('client') == "Cipla Pharmaceuticals":
                    p = add_justified_paragraph("Responsibilities: To modernize Cipla's healthcare data infrastructure", bold=True)
                else:
                    p = add_justified_paragraph("Responsibilities:", bold=True)
                add_bullet_points(job['responsibilities'])
                
            if 'environment' in job and job['environment'] and len(job['environment']) > 0:
                p = add_justified_paragraph("Environment:", bold=True)
                env_text = ", ".join(job['environment'])
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Environment section: No bold skills highlighting, just plain text
                run = p.add_run(env_text)
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Add minimal spacing between clients
            if i < len(data['experience']) - 1:
                spacing_para = doc.add_paragraph()
                spacing_para.paragraph_format.space_after = Pt(2)
    
    # --- Technical Skills ---
    if 'technical_skills' in data and data['technical_skills']:
        add_section_heading_conditional("Technical Skills")
        first_category = True
        for category, skills in data['technical_skills'].items():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Remove extra spacing for first category
            if first_category:
                p.paragraph_format.space_before = Pt(0)
                first_category = False
            run = p.add_run(f"{category}: ")
            run.font.size = Pt(11)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            
            skills_text = ", ".join(skills)
            skills_run = p.add_run(skills_text)
            skills_run.font.size = Pt(11)
            skills_run.bold = False
    
    # --- Education ---
    if 'education' in data:
        add_section_heading_conditional("Education")
        educations = data['education'] if isinstance(data['education'], list) else [data['education']]
        for edu in educations:
            p = doc.add_paragraph()
            if 'degree' in edu:
                degree_run = p.add_run(f"{edu['degree']} ")
                degree_run.bold = True
                degree_run.font.color.rgb = RGBColor(0, 0, 0)
                remaining_parts = []
                if 'field' in edu:
                    remaining_parts.append(f"in {edu['field']}")
                if 'concentration' in edu and edu['concentration']:
                    remaining_parts.append(f"({edu['concentration']})")
                if 'institution' in edu:
                    remaining_parts.append(f"| {edu['institution']}")
                if 'year' in edu and edu['year'] and str(edu['year']).strip().upper() != 'YYYY':
                    remaining_parts.append(f"({edu['year']})")
                remaining_text = " ".join(remaining_parts)
                p.add_run(remaining_text)
    
    # --- Certifications ---
    if 'certifications' in data and data['certifications'] and len(data['certifications']) > 0:
        add_section_heading_conditional("Certifications")
        # Don't bold keywords in certifications
        for item in data['certifications']:
            p = doc.add_paragraph(style='List Bullet')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(item)
            run.font.size = Pt(11)
    
    doc.save(output_path)
    return output_path


def generate_resume_style_4(json_file_path: str, resume_directory: str, font_style: str = "Calibri", list_of_keywords: Optional[List[str]] = None) -> str:
    """
    Generate resume from JSON with style_4 format: Personal Information, Professional Summary, 
    Technical Skills, Professional Experience, Education, Certifications
    WITH SPECIFIC ORDER: Summary → Skills → Experience → Education → Certifications
    
    Args:
        json_file_path: Path to JSON file with resume data
        resume_directory: Directory to save the generated resume
        font_style: Font name to use for the resume (default: "Calibri")
        list_of_keywords: List of keywords to highlight in the resume (optional)
    
    Returns:
        Path to generated resume file
    """
    # Load JSON data
    with open(json_file_path, 'r') as file:
        data = json.load(file)
    name = data.get('name', 'resume').replace(" ", "_")
    # Add timestamp to prevent file overriding
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # Create today's date folder (format: YYYY-MM-DD)
    today_folder = datetime.now().strftime("%Y-%m-%d")
    today_resume_directory = os.path.join(resume_directory, today_folder)
    
    # Ensure the today's folder exists
    os.makedirs(today_resume_directory, exist_ok=True)
    
    output_path = os.path.join(today_resume_directory, f"{name}_style_4_ordered_{timestamp}.docx")
        
    doc = Document()
    # Set tight margins (0.5 inches)
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    # Add page border (0.2 inches from page edges)
    sectPr = section._sectPr
    pgBorders = parse_xml(r'<w:pgBorders {0} w:offsetFrom="page"><w:top w:val="single" w:sz="6" w:space="14.4" w:color="000000"/><w:left w:val="single" w:sz="6" w:space="14.4" w:color="000000"/><w:bottom w:val="single" w:sz="6" w:space="14.4" w:color="000000"/><w:right w:val="single" w:sz="6" w:space="14.4" w:color="000000"/></w:pgBorders>'.format(nsdecls('w')))
    sectPr.append(pgBorders)
    
    # Configure base style with provided font
    style = doc.styles['Normal']
    font = style.font  # type: ignore
    font.name = font_style
    font.size = Pt(11)
    style.paragraph_format.space_before = Pt(0)  # type: ignore
    style.paragraph_format.space_after = Pt(0)  # type: ignore
    style.paragraph_format.line_spacing = 1.0  # type: ignore
    # Load bold keywords from bold_words.json
    bold_keywords = load_bold_keywords()
    
    # Extract all skills for highlighting (only if list_of_keywords is provided)
    all_skills = []
    if list_of_keywords:
        # Add technical skills from the resume
        if 'technical_skills' in data and data['technical_skills']:
            for category, skills in data['technical_skills'].items():
                all_skills.extend(skills)
        
        # Add provided keywords to skills list
        all_skills.extend(list_of_keywords)
    
    # Always add bold keywords to the skills list for bolding
    all_skills.extend(bold_keywords)
    # Helper function to add hyperlinks with blue color
    def add_hyperlink(paragraph, text, url):
        """Add a hyperlink to a paragraph with blue color"""
        # Create hyperlink element
        hyperlink = parse_xml(r'<w:hyperlink {0} w:history="1"><w:r><w:rPr><w:rStyle w:val="Hyperlink"/></w:rPr><w:t>{1}</w:t></w:r></w:hyperlink>'.format(nsdecls('w'), text))
        hyperlink.set(qn('r:id'), paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True))
        paragraph._p.append(hyperlink)
        
        # Set the hyperlink color to blue
        for run in hyperlink.iter():
            if run.tag.endswith('}r'):
                rPr = run.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
                if rPr is not None:
                    # Add color element for blue
                    color_elem = parse_xml(r'<w:color {0} w:val="0000FF"/>'.format(nsdecls('w')))
                    rPr.append(color_elem)
                    # Add underline element
                    underline_elem = parse_xml(r'<w:u {0} w:val="single"/>'.format(nsdecls('w')))
                    rPr.append(underline_elem)
        
        return hyperlink
    # Helper functions
    def add_centered_paragraph(text, bold=False, size=11, color=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        return p
    def add_justified_paragraph(text=None, bold=False, color=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if text:
            run = p.add_run(text)
            run.bold = bold
            if color:
                run.font.color.rgb = RGBColor(*color)
        return p
    def add_section_heading_conditional(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text.upper())
        run.bold = True
        run.font.name = font_style
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Apply light blue shading to ALL sections (remove separation line)
        shading = parse_xml(r'<w:shd {0} w:val="clear" w:color="auto" w:fill="ADD8E6"/>'.format(nsdecls('w')))
        p._p.get_or_add_pPr().append(shading)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
    def add_bullet_points(items):
        for item in items:
            p = doc.add_paragraph(style='List Bullet')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            if all_skills:
                highlight_skills_in_text(item, p, all_skills)
            else:
                run = p.add_run(item)
                run.font.size = Pt(11)
    def highlight_skills_in_text(text, paragraph, skills_list):
        """Highlight skills in text with bold formatting using exact keyword matching"""
        if not skills_list:
            run = paragraph.add_run(text)
            run.font.size = Pt(11)
            return
        
        # Create exact match patterns with word boundaries
        exact_patterns = []
        for skill in skills_list:
            # Escape special regex characters
            escaped_skill = re.escape(skill.strip())
            # Create pattern with word boundaries for exact matching
            # This ensures the skill is matched as a complete word/phrase
            pattern = r'\b' + escaped_skill + r'\b'
            exact_patterns.append((pattern, skill.strip()))
        
        # Find all matches with their positions
        matches = []
        for pattern, original_skill in exact_patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                matches.append((match.start(), match.end(), match.group(), original_skill))
        
        # Sort matches by position
        matches.sort(key=lambda x: x[0])
        
        # Remove overlapping matches (keep the first one)
        filtered_matches = []
        for match in matches:
            if not any(existing[0] <= match[0] < existing[1] or existing[0] < match[1] <= existing[1] 
                      for existing in filtered_matches):
                filtered_matches.append(match)
        
        if not filtered_matches:
            run = paragraph.add_run(text)
            run.font.size = Pt(11)
            return
        
        paragraph.clear()
        last_end = 0
        
        for start, end, matched_text, original_skill in filtered_matches:
            # Add text before the match
            if start > last_end:
                run = paragraph.add_run(text[last_end:start])
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Add the highlighted skill
            skill_run = paragraph.add_run(matched_text)
            skill_run.bold = True
            skill_run.font.color.rgb = RGBColor(0, 0, 0)
            skill_run.font.size = Pt(11)
            
            last_end = end
        
        # Add remaining text after the last match
        if last_end < len(text):
            run = paragraph.add_run(text[last_end:])
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)
    
    # --- Personal Information (Name, Title, Contact) ---
    add_centered_paragraph(data['name'], bold=True, size=16)
    add_centered_paragraph(data['title'], size=12)
    contact = data['contact']
    # Create contact paragraph with hyperlinks
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    first_item = True
    
    # Email with mailto: hyperlink
    if 'email' in contact:
        if not first_item:
            contact_p.add_run(" | ")
        add_hyperlink(contact_p, contact['email'], f"mailto:{contact['email']}")
        first_item = False
    # Phone with tel: hyperlink
    if 'phone' in contact:
        if not first_item:
            contact_p.add_run(" | ")
        add_hyperlink(contact_p, contact['phone'], f"tel:{contact['phone']}")
        first_item = False
    # LinkedIn with URL hyperlink
    if 'linkedin' in contact and contact['linkedin'] and contact['linkedin'].strip():
        if not first_item:
            contact_p.add_run(" | ")
        add_hyperlink(contact_p, "LinkedIn", contact['linkedin'])
        first_item = False
    # Portfolio with URL hyperlink
    if contact.get('portfolio'):
        if not first_item:
            contact_p.add_run(" | ")
        add_hyperlink(contact_p, "Portfolio", contact['portfolio'])
    
    # --- Professional Summary (FIRST) ---
    if 'professional_summary' in data and data['professional_summary']:
        add_section_heading_conditional("Professional Summary")
        add_bullet_points(data['professional_summary'])
    
    # --- Technical Skills (SECOND) ---
    if 'technical_skills' in data and data['technical_skills']:
        add_section_heading_conditional("Technical Skills")
        first_category = True
        for category, skills in data['technical_skills'].items():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Remove extra spacing for first category
            if first_category:
                p.paragraph_format.space_before = Pt(0)
                first_category = False
            run = p.add_run(f"{category}: ")
            run.font.size = Pt(11)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            
            skills_text = ", ".join(skills)
            skills_run = p.add_run(skills_text)
            skills_run.font.size = Pt(11)
            skills_run.bold = False
    
    # --- Professional Experience (THIRD) ---
    if 'experience' in data and data['experience']:
        add_section_heading_conditional("Professional Experience")
        for i, job in enumerate(data['experience']):
            # Job title, company, and duration in single line format
            job_para = doc.add_paragraph()
            job_para.paragraph_format.space_after = Pt(2)
            
            # Format: "Role || Client || Duration"
            role = job.get('role', 'N/A')
            company = job.get('client', job.get('company', 'N/A'))
            duration = job.get('duration', 'N/A')
            
            job_run = job_para.add_run(f"{role} || {company} || {duration}")
            job_run.bold = True
            job_run.font.color.rgb = RGBColor(0, 0, 0)
            job_run.font.size = Pt(11)
                
            if 'responsibilities' in job and job['responsibilities']:
                if job.get('client') == "Cipla Pharmaceuticals":
                    p = add_justified_paragraph("Responsibilities: To modernize Cipla's healthcare data infrastructure", bold=True)
                else:
                    p = add_justified_paragraph("Responsibilities:", bold=True)
                add_bullet_points(job['responsibilities'])
                
            if 'environment' in job and job['environment'] and len(job['environment']) > 0:
                p = add_justified_paragraph("Environment:", bold=True)
                env_text = ", ".join(job['environment'])
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Environment section: No bold skills highlighting, just plain text
                run = p.add_run(env_text)
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Add minimal spacing between clients
            if i < len(data['experience']) - 1:
                spacing_para = doc.add_paragraph()
                spacing_para.paragraph_format.space_after = Pt(2)
    
    # --- Education (FOURTH) ---
    if 'education' in data:
        add_section_heading_conditional("Education")
        educations = data['education'] if isinstance(data['education'], list) else [data['education']]
        for edu in educations:
            p = doc.add_paragraph()
            if 'degree' in edu:
                degree_run = p.add_run(f"{edu['degree']} ")
                degree_run.bold = True
                degree_run.font.color.rgb = RGBColor(0, 0, 0)
                remaining_parts = []
                if 'field' in edu:
                    remaining_parts.append(f"in {edu['field']}")
                if 'concentration' in edu and edu['concentration']:
                    remaining_parts.append(f"({edu['concentration']})")
                if 'institution' in edu:
                    remaining_parts.append(f"| {edu['institution']}")
                if 'year' in edu and edu['year'] and str(edu['year']).strip().upper() != 'YYYY':
                    remaining_parts.append(f"({edu['year']})")
                remaining_text = " ".join(remaining_parts)
                p.add_run(remaining_text)
    
    # --- Certifications (FIFTH) ---
    if 'certifications' in data and data['certifications'] and len(data['certifications']) > 0:
        add_section_heading_conditional("Certifications")
        # Don't bold keywords in certifications
        for item in data['certifications']:
            p = doc.add_paragraph(style='List Bullet')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(item)
            run.font.size = Pt(11)
    
    doc.save(output_path)
    return output_path


def generate_resume_style_5(json_file_path: str, resume_directory: str, font_style: str = "Calibri", list_of_keywords: Optional[List[str]] = None) -> str:
    """
    Generate resume from JSON with style_5 format: Personal Information, Professional Summary, 
    Technical Skills, Professional Experience, Education, Certifications
    WITH BORDER LINES and Technical Skills in second position
    
    Args:
        json_file_path: Path to JSON file with resume data
        resume_directory: Directory to save the generated resume
        font_style: Font name to use for the resume (default: "Calibri")
        list_of_keywords: List of keywords to highlight in the resume (optional)
    
    Returns:
        Path to generated resume file
    """
    # Load JSON data
    with open(json_file_path, 'r') as file:
        data = json.load(file)
    
    # Handle case where JSON contains a string instead of dict
    if isinstance(data, str):
        data = json.loads(data)
    
    name = "Yallaiah_Onteru"
    
    # Create today's date folder (format: YYYY-MM-DD)
    today_folder = datetime.now().strftime("%Y-%m-%d")
    today_resume_directory = os.path.join(resume_directory, today_folder)
    
    # Ensure the today's folder exists
    os.makedirs(today_resume_directory, exist_ok=True)
    
    output_path = os.path.join(today_resume_directory, f"{name}.docx")
        
    doc = Document()
    # Set tight margins (0.5 inches)
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    # Add page border (0.2 inches from page edges)
    sectPr = section._sectPr
    pgBorders = parse_xml(r'<w:pgBorders {0} w:offsetFrom="page"><w:top w:val="single" w:sz="6" w:space="14.4" w:color="000000"/><w:left w:val="single" w:sz="6" w:space="14.4" w:color="000000"/><w:bottom w:val="single" w:sz="6" w:space="14.4" w:color="000000"/><w:right w:val="single" w:sz="6" w:space="14.4" w:color="000000"/></w:pgBorders>'.format(nsdecls('w')))
    sectPr.append(pgBorders)
    
    # Configure base style with provided font
    style = doc.styles['Normal']
    font = style.font  # type: ignore
    font.name = font_style
    font.size = Pt(11)
    style.paragraph_format.space_before = Pt(0)  # type: ignore
    style.paragraph_format.space_after = Pt(0)  # type: ignore
    style.paragraph_format.line_spacing = 1.0  # type: ignore
    # Load bold keywords from bold_words.json
    bold_keywords = load_bold_keywords()
    
    # Extract all skills for highlighting (only if list_of_keywords is provided)
    all_skills = []
    if list_of_keywords:
        # Add technical skills from the resume
        if 'technical_skills' in data and data['technical_skills']:
            for category, skills in data['technical_skills'].items():
                all_skills.extend(skills)
        
        # Add provided keywords to skills list
        all_skills.extend(list_of_keywords)
    
    # Always add bold keywords to the skills list for bolding
    all_skills.extend(bold_keywords)
    # Helper function to add hyperlinks with blue color
    def add_hyperlink(paragraph, text, url):
        """Add a hyperlink to a paragraph with blue color"""
        # Create hyperlink element
        hyperlink = parse_xml(r'<w:hyperlink {0} w:history="1"><w:r><w:rPr><w:rStyle w:val="Hyperlink"/></w:rPr><w:t>{1}</w:t></w:r></w:hyperlink>'.format(nsdecls('w'), text))
        hyperlink.set(qn('r:id'), paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True))
        paragraph._p.append(hyperlink)
        
        # Set the hyperlink color to blue
        for run in hyperlink.iter():
            if run.tag.endswith('}r'):
                rPr = run.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
                if rPr is not None:
                    # Add color element for blue
                    color_elem = parse_xml(r'<w:color {0} w:val="0000FF"/>'.format(nsdecls('w')))
                    rPr.append(color_elem)
                    # Add underline element
                    underline_elem = parse_xml(r'<w:u {0} w:val="single"/>'.format(nsdecls('w')))
                    rPr.append(underline_elem)
        
        return hyperlink
    # Helper functions
    def add_centered_paragraph(text, bold=False, size=11, color=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        return p
    def add_justified_paragraph(text=None, bold=False, color=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if text:
            run = p.add_run(text)
            run.bold = bold
            if color:
                run.font.color.rgb = RGBColor(*color)
        return p
    def add_section_heading(text):
        # Add bold section heading with border line (no gap)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text.upper())
        run.bold = True
        run.font.name = font_style
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add border line directly under text (no gap)
        p_border = parse_xml(r'<w:pBdr {0}><w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/></w:pBdr>'.format(nsdecls('w')))
        p._p.get_or_add_pPr().append(p_border)
        
        p.paragraph_format.space_after = Pt(4)
    def add_bullet_points(items):
        for item in items:
            p = doc.add_paragraph(style='List Bullet')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            if all_skills:
                highlight_skills_in_text(item, p, all_skills)
            else:
                run = p.add_run(item)
                run.font.size = Pt(11)
    def highlight_skills_in_text(text, paragraph, skills_list):
        """Highlight skills in text with bold formatting using exact keyword matching"""
        if not skills_list:
            run = paragraph.add_run(text)
            run.font.size = Pt(11)
            return
        
        # Create exact match patterns with word boundaries
        exact_patterns = []
        for skill in skills_list:
            # Escape special regex characters
            escaped_skill = re.escape(skill.strip())
            # Create pattern with word boundaries for exact matching
            # This ensures the skill is matched as a complete word/phrase
            pattern = r'\b' + escaped_skill + r'\b'
            exact_patterns.append((pattern, skill.strip()))
        
        # Find all matches with their positions
        matches = []
        for pattern, original_skill in exact_patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                matches.append((match.start(), match.end(), match.group(), original_skill))
        
        # Sort matches by position
        matches.sort(key=lambda x: x[0])
        
        # Remove overlapping matches (keep the first one)
        filtered_matches = []
        for match in matches:
            if not any(existing[0] <= match[0] < existing[1] or existing[0] < match[1] <= existing[1] 
                      for existing in filtered_matches):
                filtered_matches.append(match)
        
        if not filtered_matches:
            run = paragraph.add_run(text)
            run.font.size = Pt(11)
            return
        
        paragraph.clear()
        last_end = 0
        
        for start, end, matched_text, original_skill in filtered_matches:
            # Add text before the match
            if start > last_end:
                run = paragraph.add_run(text[last_end:start])
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Add the highlighted skill
            skill_run = paragraph.add_run(matched_text)
            skill_run.bold = True
            skill_run.font.color.rgb = RGBColor(0, 0, 0)
            skill_run.font.size = Pt(11)
            
            last_end = end
        
        # Add remaining text after the last match
        if last_end < len(text):
            run = paragraph.add_run(text[last_end:])
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)
    
    # --- Personal Information (Name, Title, Contact) ---
    add_centered_paragraph(data['name'], bold=True, size=16)
    add_centered_paragraph(data['title'], size=12)
    contact = data['contact']
    # Create contact paragraph with hyperlinks
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    first_item = True
    
    # Email with mailto: hyperlink
    if 'email' in contact:
        if not first_item:
            contact_p.add_run(" | ")
        add_hyperlink(contact_p, contact['email'], f"mailto:{contact['email']}")
        first_item = False
    # Phone with tel: hyperlink
    if 'phone' in contact:
        if not first_item:
            contact_p.add_run(" | ")
        add_hyperlink(contact_p, contact['phone'], f"tel:{contact['phone']}")
        first_item = False
    # LinkedIn with URL hyperlink
    if 'linkedin' in contact and contact['linkedin'] and contact['linkedin'].strip():
        if not first_item:
            contact_p.add_run(" | ")
        add_hyperlink(contact_p, "LinkedIn", contact['linkedin'])
        first_item = False
    # Portfolio with URL hyperlink
    if contact.get('portfolio'):
        if not first_item:
            contact_p.add_run(" | ")
        add_hyperlink(contact_p, "Portfolio", contact['portfolio'])
    
    # --- Professional Summary (FIRST) ---
    if 'professional_summary' in data and data['professional_summary']:
        add_section_heading("Professional Summary")
        add_bullet_points(data['professional_summary'])
    
    # --- Technical Skills (SECOND) ---
    if 'technical_skills' in data and data['technical_skills']:
        add_section_heading("Technical Skills")
        first_category = True
        for category, skills in data['technical_skills'].items():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Remove extra spacing for first category
            if first_category:
                p.paragraph_format.space_before = Pt(0)
                first_category = False
            run = p.add_run(f"{category}: ")
            run.font.size = Pt(11)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            
            skills_text = ", ".join(skills)
            skills_run = p.add_run(skills_text)
            skills_run.font.size = Pt(11)
            skills_run.bold = False
    
    # --- Professional Experience (THIRD) ---
    if 'experience' in data and data['experience']:
        add_section_heading("Professional Experience")
        for i, job in enumerate(data['experience']):
            # Create paragraph for Job Title with Duration right-aligned
            role_para = doc.add_paragraph()
            role_para.paragraph_format.space_after = Pt(0)
            
            # Add job title (left-aligned, bold)
            title_run = role_para.add_run(job['role'])
            title_run.bold = True
            title_run.font.color.rgb = RGBColor(0, 0, 0)
            title_run.font.size = Pt(11)
            
            # Add tab stops for right alignment of duration
            tab_stops = role_para.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
            
            # Add tab and duration (right-aligned, bold)
            duration_run = role_para.add_run("\t" + job['duration'])
            duration_run.bold = True
            duration_run.font.color.rgb = RGBColor(0, 0, 0)
            duration_run.font.size = Pt(11)
            
            # Client and location information (next line, left-aligned)
            client = job.get('client', '').strip()
            location = job.get('location', '').strip() if 'location' in job else ''
            if client and location:
                client_info = f"{client} – {location}"
            elif client:
                client_info = client
            elif location:
                client_info = location
            else:
                client_info = ''
                
            if client_info:
                client_para = doc.add_paragraph()
                client_para.paragraph_format.space_before = Pt(0)
                client_para.paragraph_format.space_after = Pt(2)
                client_run = client_para.add_run(client_info)
                client_run.bold = True
                client_run.font.color.rgb = RGBColor(0, 0, 0)
                client_run.font.size = Pt(11)
                
            if 'responsibilities' in job and job['responsibilities']:
                if job.get('client') == "Cipla Pharmaceuticals":
                    p = add_justified_paragraph("Responsibilities: To modernize Cipla's healthcare data infrastructure", bold=True)
                else:
                    p = add_justified_paragraph("Responsibilities:", bold=True)
                add_bullet_points(job['responsibilities'])
                
            if 'environment' in job and job['environment'] and len(job['environment']) > 0:
                p = add_justified_paragraph("Environment:", bold=True)
                env_text = ", ".join(job['environment'])
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Environment section: No bold skills highlighting, just plain text
                run = p.add_run(env_text)
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Add minimal spacing between clients
            if i < len(data['experience']) - 1:
                spacing_para = doc.add_paragraph()
                spacing_para.paragraph_format.space_after = Pt(2)
    
    # --- Education (FOURTH) ---
    if 'education' in data:
        add_section_heading("Education")
        educations = data['education'] if isinstance(data['education'], list) else [data['education']]
        for edu in educations:
            p = doc.add_paragraph()
            if 'degree' in edu:
                degree_run = p.add_run(f"{edu['degree']} ")
                degree_run.bold = True
                degree_run.font.color.rgb = RGBColor(0, 0, 0)
                remaining_parts = []
                if 'field' in edu:
                    remaining_parts.append(f"in {edu['field']}")
                if 'concentration' in edu and edu['concentration']:
                    remaining_parts.append(f"({edu['concentration']})")
                if 'institution' in edu:
                    remaining_parts.append(f"| {edu['institution']}")
                if 'year' in edu and edu['year'] and str(edu['year']).strip().upper() != 'YYYY':
                    remaining_parts.append(f"({edu['year']})")
                remaining_text = " ".join(remaining_parts)
                p.add_run(remaining_text)
    
    # --- Certifications (FIFTH) ---
    if 'certifications' in data and data['certifications'] and len(data['certifications']) > 0:
        add_section_heading("Certifications")
        # Don't bold keywords in certifications
        for item in data['certifications']:
            p = doc.add_paragraph(style='List Bullet')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(item)
            run.font.size = Pt(11)
    
    doc.save(output_path)
    return output_path


def generate_resume_style_6(json_file_path: str, resume_directory: str, font_style: str = "Calibri", list_of_keywords: Optional[List[str]] = None) -> str:
    """
    Generate resume from JSON with style_6 format: Personal Information, Professional Summary, 
    Technical Skills, Professional Experience, Education, Certifications
    WITH BORDER LINES and Technical Skills in second position
    
    Args:
        json_file_path: Path to JSON file with resume data
        resume_directory: Directory to save the generated resume
        font_style: Font name to use for the resume (default: "Calibri")
        list_of_keywords: List of keywords to highlight in the resume (optional)
    
    Returns:
        Path to generated resume file
    """
    # Load JSON data
    with open(json_file_path, 'r') as file:
        data = json.load(file)
    name = data.get('name', 'resume').replace(" ", "_")
    # Add timestamp to prevent file overriding
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # Create today's date folder (format: YYYY-MM-DD)
    today_folder = datetime.now().strftime("%Y-%m-%d")
    today_resume_directory = os.path.join(resume_directory, today_folder)
    
    # Ensure the today's folder exists
    os.makedirs(today_resume_directory, exist_ok=True)
    
    output_path = os.path.join(today_resume_directory, f"{name}_style_6_shading_tech_second_{timestamp}.docx")
        
    doc = Document()
    # Set tight margins (0.5 inches)
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    # Add page border (0.2 inches from page edges)
    sectPr = section._sectPr
    pgBorders = parse_xml(r'<w:pgBorders {0} w:offsetFrom="page"><w:top w:val="single" w:sz="6" w:space="14.4" w:color="000000"/><w:left w:val="single" w:sz="6" w:space="14.4" w:color="000000"/><w:bottom w:val="single" w:sz="6" w:space="14.4" w:color="000000"/><w:right w:val="single" w:sz="6" w:space="14.4" w:color="000000"/></w:pgBorders>'.format(nsdecls('w')))
    sectPr.append(pgBorders)
    
    # Configure base style with provided font
    style = doc.styles['Normal']
    font = style.font  # type: ignore
    font.name = font_style
    font.size = Pt(11)
    style.paragraph_format.space_before = Pt(0)  # type: ignore
    style.paragraph_format.space_after = Pt(0)  # type: ignore
    style.paragraph_format.line_spacing = 1.0  # type: ignore
    # Load bold keywords from bold_words.json
    bold_keywords = load_bold_keywords()
    
    # Extract all skills for highlighting (only if list_of_keywords is provided)
    all_skills = []
    if list_of_keywords:
        # Add technical skills from the resume
        if 'technical_skills' in data and data['technical_skills']:
            for category, skills in data['technical_skills'].items():
                all_skills.extend(skills)
        
        # Add provided keywords to skills list
        all_skills.extend(list_of_keywords)
    
    # Always add bold keywords to the skills list for bolding
    all_skills.extend(bold_keywords)
    # Helper function to add hyperlinks with blue color
    def add_hyperlink(paragraph, text, url):
        """Add a hyperlink to a paragraph with blue color"""
        # Create hyperlink element
        hyperlink = parse_xml(r'<w:hyperlink {0} w:history="1"><w:r><w:rPr><w:rStyle w:val="Hyperlink"/></w:rPr><w:t>{1}</w:t></w:r></w:hyperlink>'.format(nsdecls('w'), text))
        hyperlink.set(qn('r:id'), paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True))
        paragraph._p.append(hyperlink)
        
        # Set the hyperlink color to blue
        for run in hyperlink.iter():
            if run.tag.endswith('}r'):
                rPr = run.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
                if rPr is not None:
                    # Add color element for blue
                    color_elem = parse_xml(r'<w:color {0} w:val="0000FF"/>'.format(nsdecls('w')))
                    rPr.append(color_elem)
                    # Add underline element
                    underline_elem = parse_xml(r'<w:u {0} w:val="single"/>'.format(nsdecls('w')))
                    rPr.append(underline_elem)
        
        return hyperlink
    # Helper functions
    def add_centered_paragraph(text, bold=False, size=11, color=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        return p
    def add_justified_paragraph(text=None, bold=False, color=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if text:
            run = p.add_run(text)
            run.bold = bold
            if color:
                run.font.color.rgb = RGBColor(*color)
        return p
    def add_section_heading(text):
        # Add bold section heading with border line (no gap)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text.upper())
        run.bold = True
        run.font.name = font_style
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add border line directly under text (no gap)
        p_border = parse_xml(r'<w:pBdr {0}><w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/></w:pBdr>'.format(nsdecls('w')))
        p._p.get_or_add_pPr().append(p_border)
        
        p.paragraph_format.space_after = Pt(4)
    def add_bullet_points(items):
        for item in items:
            p = doc.add_paragraph(style='List Bullet')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            if all_skills:
                highlight_skills_in_text(item, p, all_skills)
            else:
                run = p.add_run(item)
                run.font.size = Pt(11)
    def highlight_skills_in_text(text, paragraph, skills_list):
        """Highlight skills in text with bold formatting using exact keyword matching"""
        if not skills_list:
            run = paragraph.add_run(text)
            run.font.size = Pt(11)
            return
        
        # Create exact match patterns with word boundaries
        exact_patterns = []
        for skill in skills_list:
            # Escape special regex characters
            escaped_skill = re.escape(skill.strip())
            # Create pattern with word boundaries for exact matching
            # This ensures the skill is matched as a complete word/phrase
            pattern = r'\b' + escaped_skill + r'\b'
            exact_patterns.append((pattern, skill.strip()))
        
        # Find all matches with their positions
        matches = []
        for pattern, original_skill in exact_patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                matches.append((match.start(), match.end(), match.group(), original_skill))
        
        # Sort matches by position
        matches.sort(key=lambda x: x[0])
        
        # Remove overlapping matches (keep the first one)
        filtered_matches = []
        for match in matches:
            if not any(existing[0] <= match[0] < existing[1] or existing[0] < match[1] <= existing[1] 
                      for existing in filtered_matches):
                filtered_matches.append(match)
        
        if not filtered_matches:
            run = paragraph.add_run(text)
            run.font.size = Pt(11)
            return
        
        paragraph.clear()
        last_end = 0
        
        for start, end, matched_text, original_skill in filtered_matches:
            # Add text before the match
            if start > last_end:
                run = paragraph.add_run(text[last_end:start])
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Add the highlighted skill
            skill_run = paragraph.add_run(matched_text)
            skill_run.bold = True
            skill_run.font.color.rgb = RGBColor(0, 0, 0)
            skill_run.font.size = Pt(11)
            
            last_end = end
        
        # Add remaining text after the last match
        if last_end < len(text):
            run = paragraph.add_run(text[last_end:])
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)
    
    # --- Personal Information (Name, Title, Contact) ---
    add_centered_paragraph(data['name'], bold=True, size=16)
    add_centered_paragraph(data['title'], size=12)
    contact = data['contact']
    # Create contact paragraph with hyperlinks
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    first_item = True
    
    # Email with mailto: hyperlink
    if 'email' in contact:
        if not first_item:
            contact_p.add_run(" | ")
        add_hyperlink(contact_p, contact['email'], f"mailto:{contact['email']}")
        first_item = False
    # Phone with tel: hyperlink
    if 'phone' in contact:
        if not first_item:
            contact_p.add_run(" | ")
        add_hyperlink(contact_p, contact['phone'], f"tel:{contact['phone']}")
        first_item = False
    # LinkedIn with URL hyperlink
    if 'linkedin' in contact and contact['linkedin'] and contact['linkedin'].strip():
        if not first_item:
            contact_p.add_run(" | ")
        add_hyperlink(contact_p, "LinkedIn", contact['linkedin'])
        first_item = False
    # Portfolio with URL hyperlink
    if contact.get('portfolio'):
        if not first_item:
            contact_p.add_run(" | ")
        add_hyperlink(contact_p, "Portfolio", contact['portfolio'])
    
    # --- Professional Summary (FIRST) ---
    if 'professional_summary' in data and data['professional_summary']:
        add_section_heading("Professional Summary")
        add_bullet_points(data['professional_summary'])
    
    # --- Technical Skills (SECOND) ---
    if 'technical_skills' in data and data['technical_skills']:
        add_section_heading("Technical Skills")
        first_category = True
        for category, skills in data['technical_skills'].items():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Remove extra spacing for first category
            if first_category:
                p.paragraph_format.space_before = Pt(0)
                first_category = False
            run = p.add_run(f"{category}: ")
            run.font.size = Pt(11)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            
            skills_text = ", ".join(skills)
            skills_run = p.add_run(skills_text)
            skills_run.font.size = Pt(11)
            skills_run.bold = False
    
    # --- Professional Experience (THIRD) ---
    if 'experience' in data and data['experience']:
        add_section_heading("Professional Experience")
        for i, job in enumerate(data['experience']):
            # Create paragraph for Job Title with Duration right-aligned
            role_para = doc.add_paragraph()
            role_para.paragraph_format.space_after = Pt(0)
            
            # Add job title (left-aligned, bold)
            title_run = role_para.add_run(job['role'])
            title_run.bold = True
            title_run.font.color.rgb = RGBColor(0, 0, 0)
            title_run.font.size = Pt(11)
            
            # Add tab stops for right alignment of duration
            tab_stops = role_para.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
            
            # Add tab and duration (right-aligned, bold)
            duration_run = role_para.add_run("\t" + job['duration'])
            duration_run.bold = True
            duration_run.font.color.rgb = RGBColor(0, 0, 0)
            duration_run.font.size = Pt(11)
            
            # Client and location information (next line, left-aligned)
            client = job.get('client', '').strip()
            location = job.get('location', '').strip() if 'location' in job else ''
            if client and location:
                client_info = f"{client} – {location}"
            elif client:
                client_info = client
            elif location:
                client_info = location
            else:
                client_info = ''
                
            if client_info:
                client_para = doc.add_paragraph()
                client_para.paragraph_format.space_before = Pt(0)
                client_para.paragraph_format.space_after = Pt(2)
                client_run = client_para.add_run(client_info)
                client_run.bold = True
                client_run.font.color.rgb = RGBColor(0, 0, 0)
                client_run.font.size = Pt(11)
                
            if 'responsibilities' in job and job['responsibilities']:
                if job.get('client') == "Cipla Pharmaceuticals":
                    p = add_justified_paragraph("Responsibilities: To modernize Cipla's healthcare data infrastructure", bold=True)
                else:
                    p = add_justified_paragraph("Responsibilities:", bold=True)
                add_bullet_points(job['responsibilities'])
                
            if 'environment' in job and job['environment'] and len(job['environment']) > 0:
                p = add_justified_paragraph("Environment:", bold=True)
                env_text = ", ".join(job['environment'])
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Environment section: No bold skills highlighting, just plain text
                run = p.add_run(env_text)
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Add minimal spacing between clients
            if i < len(data['experience']) - 1:
                spacing_para = doc.add_paragraph()
                spacing_para.paragraph_format.space_after = Pt(2)
    
    # --- Education (FOURTH) ---
    if 'education' in data:
        add_section_heading("Education")
        educations = data['education'] if isinstance(data['education'], list) else [data['education']]
        for edu in educations:
            p = doc.add_paragraph()
            if 'degree' in edu:
                degree_run = p.add_run(f"{edu['degree']} ")
                degree_run.bold = True
                degree_run.font.color.rgb = RGBColor(0, 0, 0)
                remaining_parts = []
                if 'field' in edu:
                    remaining_parts.append(f"in {edu['field']}")
                if 'concentration' in edu and edu['concentration']:
                    remaining_parts.append(f"({edu['concentration']})")
                if 'institution' in edu:
                    remaining_parts.append(f"| {edu['institution']}")
                if 'year' in edu and edu['year'] and str(edu['year']).strip().upper() != 'YYYY':
                    remaining_parts.append(f"({edu['year']})")
                remaining_text = " ".join(remaining_parts)
                p.add_run(remaining_text)
    
    # --- Certifications (FIFTH) ---
    if 'certifications' in data and data['certifications'] and len(data['certifications']) > 0:
        add_section_heading("Certifications")
        # Don't bold keywords in certifications
        for item in data['certifications']:
            p = doc.add_paragraph(style='List Bullet')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(item)
            run.font.size = Pt(11)
    
    doc.save(output_path)
    return output_path


def generate_resume_style_7(json_file_path: str, resume_directory: str, font_style: str = "Calibri", list_of_keywords: Optional[List[str]] = None) -> str:
    """
    Generate resume from JSON with style_7 format: Personal Information (LEFT ALIGNED), Professional Summary, 
    Professional Experience, Technical Skills, Education, Certifications
    WITH NAME IN LEFT CORNER and border lines
    
    Args:
        json_file_path: Path to JSON file with resume data
        resume_directory: Directory to save the generated resume
        font_style: Font name to use for the resume (default: "Calibri")
        list_of_keywords: List of keywords to highlight in the resume (optional)
    
    Returns:
        Path to generated resume file
    """
    # Load JSON data
    with open(json_file_path, 'r') as file:
        data = json.load(file)
    name = data.get('name', 'resume').replace(" ", "_")
    # Add timestamp to prevent file overriding
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # Create today's date folder (format: YYYY-MM-DD)
    today_folder = datetime.now().strftime("%Y-%m-%d")
    today_resume_directory = os.path.join(resume_directory, today_folder)
    
    # Ensure the today's folder exists
    os.makedirs(today_resume_directory, exist_ok=True)
    
    output_path = os.path.join(today_resume_directory, f"{name}_style_7_left_name_{timestamp}.docx")
        
    doc = Document()
    # Set tight margins (0.5 inches)
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    # Add page border (0.2 inches from page edges)
    sectPr = section._sectPr
    pgBorders = parse_xml(r'<w:pgBorders {0} w:offsetFrom="page"><w:top w:val="single" w:sz="6" w:space="14.4" w:color="000000"/><w:left w:val="single" w:sz="6" w:space="14.4" w:color="000000"/><w:bottom w:val="single" w:sz="6" w:space="14.4" w:color="000000"/><w:right w:val="single" w:sz="6" w:space="14.4" w:color="000000"/></w:pgBorders>'.format(nsdecls('w')))
    sectPr.append(pgBorders)
    
    # Configure base style with provided font
    style = doc.styles['Normal']
    font = style.font  # type: ignore
    font.name = font_style
    font.size = Pt(11)
    style.paragraph_format.space_before = Pt(0)  # type: ignore
    style.paragraph_format.space_after = Pt(0)  # type: ignore
    style.paragraph_format.line_spacing = 1.0  # type: ignore
    # Load bold keywords from bold_words.json
    bold_keywords = load_bold_keywords()
    
    # Extract all skills for highlighting (only if list_of_keywords is provided)
    all_skills = []
    if list_of_keywords:
        # Add technical skills from the resume
        if 'technical_skills' in data and data['technical_skills']:
            for category, skills in data['technical_skills'].items():
                all_skills.extend(skills)
        
        # Add provided keywords to skills list
        all_skills.extend(list_of_keywords)
    
    # Always add bold keywords to the skills list for bolding
    all_skills.extend(bold_keywords)
    # Helper function to add hyperlinks with blue color
    def add_hyperlink(paragraph, text, url):
        """Add a hyperlink to a paragraph with blue color"""
        # Create hyperlink element
        hyperlink = parse_xml(r'<w:hyperlink {0} w:history="1"><w:r><w:rPr><w:rStyle w:val="Hyperlink"/></w:rPr><w:t>{1}</w:t></w:r></w:hyperlink>'.format(nsdecls('w'), text))
        hyperlink.set(qn('r:id'), paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True))
        paragraph._p.append(hyperlink)
        
        # Set the hyperlink color to blue
        for run in hyperlink.iter():
            if run.tag.endswith('}r'):
                rPr = run.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
                if rPr is not None:
                    # Add color element for blue
                    color_elem = parse_xml(r'<w:color {0} w:val="0000FF"/>'.format(nsdecls('w')))
                    rPr.append(color_elem)
                    # Add underline element
                    underline_elem = parse_xml(r'<w:u {0} w:val="single"/>'.format(nsdecls('w')))
                    rPr.append(underline_elem)
        
        return hyperlink
    # Helper functions
    def add_left_paragraph(text, bold=False, size=11, color=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        return p
    def add_justified_paragraph(text=None, bold=False, color=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if text:
            run = p.add_run(text)
            run.bold = bold
            if color:
                run.font.color.rgb = RGBColor(*color)
        return p
    def add_section_heading(text):
        # Add bold section heading with border line (no gap)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text.upper())
        run.bold = True
        run.font.name = font_style
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add border line directly under text (no gap)
        p_border = parse_xml(r'<w:pBdr {0}><w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/></w:pBdr>'.format(nsdecls('w')))
        p._p.get_or_add_pPr().append(p_border)
        
        p.paragraph_format.space_after = Pt(4)
    def add_bullet_points(items):
        for item in items:
            p = doc.add_paragraph(style='List Bullet')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            if all_skills:
                highlight_skills_in_text(item, p, all_skills)
            else:
                run = p.add_run(item)
                run.font.size = Pt(11)
    def highlight_skills_in_text(text, paragraph, skills_list):
        """Highlight skills in text with bold formatting using exact keyword matching"""
        if not skills_list:
            run = paragraph.add_run(text)
            run.font.size = Pt(11)
            return
        
        # Create exact match patterns with word boundaries
        exact_patterns = []
        for skill in skills_list:
            # Escape special regex characters
            escaped_skill = re.escape(skill.strip())
            # Create pattern with word boundaries for exact matching
            # This ensures the skill is matched as a complete word/phrase
            pattern = r'\b' + escaped_skill + r'\b'
            exact_patterns.append((pattern, skill.strip()))
        
        # Find all matches with their positions
        matches = []
        for pattern, original_skill in exact_patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                matches.append((match.start(), match.end(), match.group(), original_skill))
        
        # Sort matches by position
        matches.sort(key=lambda x: x[0])
        
        # Remove overlapping matches (keep the first one)
        filtered_matches = []
        for match in matches:
            if not any(existing[0] <= match[0] < existing[1] or existing[0] < match[1] <= existing[1] 
                      for existing in filtered_matches):
                filtered_matches.append(match)
        
        if not filtered_matches:
            run = paragraph.add_run(text)
            run.font.size = Pt(11)
            return
        
        paragraph.clear()
        last_end = 0
        
        for start, end, matched_text, original_skill in filtered_matches:
            # Add text before the match
            if start > last_end:
                run = paragraph.add_run(text[last_end:start])
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Add the highlighted skill
            skill_run = paragraph.add_run(matched_text)
            skill_run.bold = True
            skill_run.font.color.rgb = RGBColor(0, 0, 0)
            skill_run.font.size = Pt(11)
            
            last_end = end
        
        # Add remaining text after the last match
        if last_end < len(text):
            run = paragraph.add_run(text[last_end:])
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)
    
    # --- Personal Information (Name, Title, Contact) - LEFT ALIGNED ---
    # Name on the left
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name_run = name_para.add_run(data.get('name', 'N/A'))
    name_run.bold = True
    name_run.font.name = font_style
    name_run.font.size = Pt(16)
    name_run.font.color.rgb = RGBColor(0, 0, 0)
    name_para.paragraph_format.space_after = Pt(2)
    
    # Title on the left
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_para.add_run(data.get('title', ''))
    title_run.bold = True
    title_run.font.name = font_style
    title_run.font.size = Pt(12)
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title_para.paragraph_format.space_after = Pt(4)
    
    # Contact information on the left - single line format
    contact = data.get('contact', {})
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    first_item = True
    
    # Email with mailto: hyperlink
    if 'email' in contact:
        if not first_item:
            contact_p.add_run(" | ")
        add_hyperlink(contact_p, contact['email'], f"mailto:{contact['email']}")
        first_item = False
    # Phone with tel: hyperlink
    if 'phone' in contact:
        if not first_item:
            contact_p.add_run(" | ")
        add_hyperlink(contact_p, contact['phone'], f"tel:{contact['phone']}")
        first_item = False
    # LinkedIn with URL hyperlink
    if 'linkedin' in contact and contact['linkedin'] and contact['linkedin'].strip():
        if not first_item:
            contact_p.add_run(" | ")
        add_hyperlink(contact_p, "LinkedIn", contact['linkedin'])
        first_item = False
    # Portfolio/GitHub with URL hyperlink
    if contact.get('portfolio'):
        if not first_item:
            contact_p.add_run(" | ")
        add_hyperlink(contact_p, "GitHub", contact['portfolio'])
        first_item = False
    
    # No space after header
    
    # --- Professional Summary ---
    if 'professional_summary' in data and data['professional_summary']:
        add_section_heading("Professional Summary")
        add_bullet_points(data['professional_summary'])
    
    # --- Professional Experience ---
    if 'experience' in data and data['experience']:
        add_section_heading("Professional Experience")
        for i, job in enumerate(data['experience']):
            # Create paragraph for Job Title with Duration right-aligned
            role_para = doc.add_paragraph()
            role_para.paragraph_format.space_after = Pt(0)
            
            # Add job title (left-aligned, bold)
            title_run = role_para.add_run(job['role'])
            title_run.bold = True
            title_run.font.color.rgb = RGBColor(0, 0, 0)
            title_run.font.size = Pt(11)
            
            # Add tab stops for right alignment of duration
            tab_stops = role_para.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
            
            # Add tab and duration (right-aligned, bold)
            duration_run = role_para.add_run("\t" + job['duration'])
            duration_run.bold = True
            duration_run.font.color.rgb = RGBColor(0, 0, 0)
            duration_run.font.size = Pt(11)
            
            # Client and location information (next line, left-aligned)
            client = job.get('client', '').strip()
            location = job.get('location', '').strip() if 'location' in job else ''
            if client and location:
                client_info = f"{client} – {location}"
            elif client:
                client_info = client
            elif location:
                client_info = location
            else:
                client_info = ''
                
            if client_info:
                client_para = doc.add_paragraph()
                client_para.paragraph_format.space_before = Pt(0)
                client_para.paragraph_format.space_after = Pt(2)
                client_run = client_para.add_run(client_info)
                client_run.bold = True
                client_run.font.color.rgb = RGBColor(0, 0, 0)
                client_run.font.size = Pt(11)
                
            if 'responsibilities' in job and job['responsibilities']:
                if job.get('client') == "Cipla Pharmaceuticals":
                    p = add_justified_paragraph("Responsibilities: To modernize Cipla's healthcare data infrastructure", bold=True)
                else:
                    p = add_justified_paragraph("Responsibilities:", bold=True)
                add_bullet_points(job['responsibilities'])
                
            if 'environment' in job and job['environment'] and len(job['environment']) > 0:
                p = add_justified_paragraph("Environment:", bold=True)
                env_text = ", ".join(job['environment'])
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Environment section: No bold skills highlighting, just plain text
                run = p.add_run(env_text)
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Add minimal spacing between clients
            if i < len(data['experience']) - 1:
                spacing_para = doc.add_paragraph()
                spacing_para.paragraph_format.space_after = Pt(2)
    
    # --- Technical Skills ---
    if 'technical_skills' in data and data['technical_skills']:
        add_section_heading("Technical Skills")
        first_category = True
        for category, skills in data['technical_skills'].items():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Remove extra spacing for first category
            if first_category:
                p.paragraph_format.space_before = Pt(0)
                first_category = False
            run = p.add_run(f"{category}: ")
            run.font.size = Pt(11)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            
            skills_text = ", ".join(skills)
            skills_run = p.add_run(skills_text)
            skills_run.font.size = Pt(11)
            skills_run.bold = False
    
    # --- Education ---
    if 'education' in data:
        add_section_heading("Education")
        educations = data['education'] if isinstance(data['education'], list) else [data['education']]
        for edu in educations:
            p = doc.add_paragraph()
            if 'degree' in edu:
                degree_run = p.add_run(f"{edu['degree']} ")
                degree_run.bold = True
                degree_run.font.color.rgb = RGBColor(0, 0, 0)
                remaining_parts = []
                if 'field' in edu:
                    remaining_parts.append(f"in {edu['field']}")
                if 'concentration' in edu and edu['concentration']:
                    remaining_parts.append(f"({edu['concentration']})")
                if 'institution' in edu:
                    remaining_parts.append(f"| {edu['institution']}")
                if 'year' in edu and edu['year'] and str(edu['year']).strip().upper() != 'YYYY':
                    remaining_parts.append(f"({edu['year']})")
                remaining_text = " ".join(remaining_parts)
                p.add_run(remaining_text)
    
    # --- Certifications ---
    if 'certifications' in data and data['certifications'] and len(data['certifications']) > 0:
        add_section_heading("Certifications")
        # Don't bold keywords in certifications
        for item in data['certifications']:
            p = doc.add_paragraph(style='List Bullet')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(item)
            run.font.size = Pt(11)
    
    doc.save(output_path)
    return output_path
#==========================================================================================================

def generate_resume_dotnet_format(json_file_path: str, resume_directory: str, font_style: str = "Times New Roman", list_of_keywords: Optional[List[str]] = None) -> str:
    """
    Generate resume from .NET JSON format with personal_info, project arrays, etc.
    
    Args:
        json_file_path: Path to JSON file with resume data (or JSON string)
        resume_directory: Directory to save the generated resume
        font_style: Font name to use for the resume (default: "Times New Roman")
        list_of_keywords: List of keywords to highlight in the resume (optional)
    
    Returns:
        Path to generated resume file
    """
    # Load JSON data - handle both file path and JSON string
    if isinstance(json_file_path, str) and os.path.exists(json_file_path):
        with open(json_file_path, 'r') as file:
            data = json.load(file)
    else:
        # Assume it's a JSON string
        data = json.loads(json_file_path)
    
    # Handle case where JSON contains a string instead of dict
    if isinstance(data, str):
        data = json.loads(data)
    
    # Extract name from personal_info
    name = "resume"
    if 'personal_info' in data and 'name' in data['personal_info']:
        name = data['personal_info']['name'].replace(" ", "_")
    elif 'name' in data:
        name = data['name'].replace(" ", "_")
    
    # Add timestamp to prevent file overriding
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # Create today's date folder
    today_folder = datetime.now().strftime("%Y-%m-%d")
    today_resume_directory = os.path.join(resume_directory, today_folder)
    os.makedirs(today_resume_directory, exist_ok=True)
    
    output_path = os.path.join(today_resume_directory, f"{name}_dotnet_{timestamp}.docx")
    
    doc = Document()
    # Set margins
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    # Configure base style
    style = doc.styles['Normal']
    font = style.font
    font.name = font_style
    font.size = Pt(9)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0
    
    # Load bold keywords from bold_words.json
    bold_keywords = load_bold_keywords()
    
    # Extract all skills for highlighting
    all_skills = []
    if list_of_keywords:
        if 'technical_skills' in data and data['technical_skills']:
            for category, skills in data['technical_skills'].items():
                if isinstance(skills, list):
                    all_skills.extend(skills)
        all_skills.extend(list_of_keywords)
    
    # Always add bold keywords
    all_skills.extend(bold_keywords)
    
    # Helper functions
    def add_hyperlink(paragraph, text, url):
        """Add a hyperlink to a paragraph with font size and blue color"""
        # Create hyperlink element with font size (18 = 9pt in half-points) and blue color
        hyperlink = parse_xml(r'<w:hyperlink {0} w:history="1"><w:r><w:rPr><w:rStyle w:val="Hyperlink"/><w:sz w:val="18"/><w:szCs w:val="18"/><w:color w:val="0000FF"/></w:rPr><w:t>{1}</w:t></w:r></w:hyperlink>'.format(nsdecls('w'), text))
        hyperlink.set(qn('r:id'), paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True))
        paragraph._p.append(hyperlink)
        return hyperlink
    
    def add_centered_paragraph(text, bold=False, size=9):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        return p
    
    def add_section_heading(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(9)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        # Add border line
        p_border = parse_xml(r'<w:pBdr {0}><w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/></w:pBdr>'.format(nsdecls('w')))
        p._p.get_or_add_pPr().append(p_border)
    
    def add_section_heading_no_border(text):
        """Add section heading without border line"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(9)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
    
    def highlight_skills_in_text(text, paragraph, skills_list):
        """Highlight skills in text with bold formatting"""
        if not skills_list or not text:
            run = paragraph.add_run(text or '')
            run.font.size = Pt(9)
            return
        
        exact_patterns = []
        for skill in skills_list:
            escaped_skill = re.escape(skill.strip())
            pattern = r'\b' + escaped_skill + r'\b'
            exact_patterns.append((pattern, skill.strip()))
        
        matches = []
        for pattern, original_skill in exact_patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                matches.append((match.start(), match.end(), match.group(), original_skill))
        
        matches.sort(key=lambda x: x[0])
        
        filtered_matches = []
        for match in matches:
            if not any(existing[0] <= match[0] < existing[1] or existing[0] < match[1] <= existing[1] 
                      for existing in filtered_matches):
                filtered_matches.append(match)
        
        if not filtered_matches:
            run = paragraph.add_run(text)
            run.font.size = Pt(9)
            return
        
        paragraph.clear()
        last_end = 0
        
        for start, end, matched_text, original_skill in filtered_matches:
            if start > last_end:
                run = paragraph.add_run(text[last_end:start])
                run.font.size = Pt(9)
            
            skill_run = paragraph.add_run(matched_text)
            skill_run.bold = True
            skill_run.font.size = Pt(9)
            
            last_end = end
        
        if last_end < len(text):
            run = paragraph.add_run(text[last_end:])
            run.font.size = Pt(9)
    
    def add_bullet_points(items):
        for item in items:
            p = doc.add_paragraph(style='List Bullet')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if all_skills:
                highlight_skills_in_text(item, p, all_skills)
            else:
                run = p.add_run(item)
                run.font.size = Pt(9)
    
    # --- Personal Information (Left Aligned) ---
    personal_info = data.get('personal_info', {})
    name_val = personal_info.get('name', '')
    title_val = personal_info.get('title', '')
    
    # Name - left aligned
    if name_val:
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        name_run = name_para.add_run(name_val)
        name_run.bold = True
        name_run.font.size = Pt(12)
        name_run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Title - left aligned
    if title_val:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_run = title_para.add_run(title_val)
        title_run.font.size = Pt(10)
        title_run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Contact info - left aligned with hyperlinks
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    first_item = True
    
    if personal_info.get('email'):
        if not first_item:
            sep_run = contact_p.add_run(" | ")
            sep_run.font.size = Pt(9)
            sep_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
        add_hyperlink(contact_p, personal_info['email'], f"mailto:{personal_info['email']}")
        first_item = False
    
    if personal_info.get('phone'):
        if not first_item:
            sep_run = contact_p.add_run(" | ")
            sep_run.font.size = Pt(9)
            sep_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
        phone_run = contact_p.add_run(personal_info['phone'])
        phone_run.font.size = Pt(9)
        phone_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
        first_item = False
    
    if personal_info.get('location'):
        if not first_item:
            sep_run = contact_p.add_run(" | ")
            sep_run.font.size = Pt(9)
            sep_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
        location_run = contact_p.add_run(personal_info['location'])
        location_run.font.size = Pt(9)
        location_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
        first_item = False
    
    if personal_info.get('linkedin'):
        if not first_item:
            sep_run = contact_p.add_run(" | ")
            sep_run.font.size = Pt(9)
            sep_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
        linkedin_url = personal_info['linkedin']
        if not linkedin_url.startswith('http'):
            linkedin_url = f"https://{linkedin_url}"
        add_hyperlink(contact_p, "LinkedIn", linkedin_url)
    
    # --- Professional Summary ---
    if 'professional_summary' in data and data['professional_summary']:
        # Add very small space before section
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run("PROFESSIONAL SUMMARY")
        run.bold = True
        run.font.size = Pt(9)
        p.paragraph_format.space_before = Pt(2)  # Very small space
        p.paragraph_format.space_after = Pt(2)
        # Add border line
        p_border = parse_xml(r'<w:pBdr {0}><w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/></w:pBdr>'.format(nsdecls('w')))
        p._p.get_or_add_pPr().append(p_border)
        add_bullet_points(data['professional_summary'])
    
    # --- Technical Skills (Table Format) ---
    if 'technical_skills' in data and data['technical_skills']:
        add_section_heading_no_border("Technical Skills")
        
        # Create plain table with 2 columns (no style to keep it simple)
        table = doc.add_table(rows=len(data['technical_skills']), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Set table column widths (30% for category, 70% for skills)
        for row in table.rows:
            row.cells[0].width = Inches(2.0)
            row.cells[1].width = Inches(5.5)
        
        row_idx = 0
        for category, skills in data['technical_skills'].items():
            # Category cell (left column) - bold, Calibri, size 9 (same as bullets)
            category_cell = table.rows[row_idx].cells[0]
            category_para = category_cell.paragraphs[0]
            category_para.clear()
            category_run = category_para.add_run(category.replace('_', ' ').title())
            category_run.font.name = 'Calibri'
            category_run.font.size = Pt(9)
            category_run.bold = True
            category_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Skills cell (right column) - Calibri, size 9 (same as bullets)
            skills_cell = table.rows[row_idx].cells[1]
            skills_para = skills_cell.paragraphs[0]
            skills_para.clear()
            
            if isinstance(skills, list):
                skills_text = ", ".join(skills)
            else:
                skills_text = str(skills)
            
            skills_run = skills_para.add_run(skills_text)
            skills_run.font.name = 'Calibri'
            skills_run.font.size = Pt(9)
            skills_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add simple black borders to cells (plain style)
            for cell in table.rows[row_idx].cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                # Simple black borders only
                borders = parse_xml(r'<w:tcBorders {0}><w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tcBorders>'.format(nsdecls('w')))
                tcPr.append(borders)
                # Remove any background shading (keep it plain)
                shading = tcPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                if shading is not None:
                    tcPr.remove(shading)
            
            row_idx += 1
    
    # --- Professional Experience ---
    if 'experience' in data and data['experience']:
        add_section_heading("Professional Experience")
        for exp in data['experience']:
            # Get company, location, role, and dates for metadata
            company = exp.get('company', '')
            location = exp.get('location', '')
            role = exp.get('role', '')
            start_date = exp.get('start_date', '')
            end_date = exp.get('end_date', '')
            
            # Projects
            if 'project' in exp and exp['project']:
                for proj in exp['project']:
                    # Add metadata section before project
                    # Client line with date right-aligned
                    client_para = doc.add_paragraph()
                    client_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Add tab stop for right alignment of date
                    tab_stops = client_para.paragraph_format.tab_stops
                    tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
                    
                    client_text = f"Client: {company}, {location}."
                    client_run = client_para.add_run(client_text)
                    client_run.bold = True
                    client_run.font.size = Pt(9)
                    client_run.font.color.rgb = RGBColor(0, 0, 0)
                    
                    # Date on right side
                    date_text = f"{start_date} - {end_date}"
                    date_run = client_para.add_run("\t" + date_text)
                    date_run.bold = True
                    date_run.font.size = Pt(9)
                    date_run.font.color.rgb = RGBColor(0, 0, 0)
                    
                    # Role line
                    role_para = doc.add_paragraph()
                    role_text = f"Role: {role}."
                    role_run = role_para.add_run(role_text)
                    role_run.bold = True
                    role_run.font.size = Pt(9)
                    role_run.font.color.rgb = RGBColor(0, 0, 0)
                    
                    # Project line
                    if proj.get('name'):
                        project_para = doc.add_paragraph()
                        project_text = f"Project: {proj['name']}."
                        project_run = project_para.add_run(project_text)
                        project_run.bold = True
                        project_run.font.size = Pt(9)
                        project_run.font.color.rgb = RGBColor(0, 0, 0)
                    
                    # Small space after metadata
                    spacing_para = doc.add_paragraph()
                    spacing_para.paragraph_format.space_after = Pt(2)
                    
                    if proj.get('project_summary'):
                        proj_summary = doc.add_paragraph()
                        proj_summary.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        if all_skills:
                            highlight_skills_in_text(proj['project_summary'], proj_summary, all_skills)
                        else:
                            proj_summary.add_run(proj['project_summary']).font.size = Pt(9)
                        proj_summary.paragraph_format.space_after = Pt(4)
                    
                    if proj.get('responsibilities'):
                        add_bullet_points(proj['responsibilities'])
                    
                    if proj.get('environment'):
                        env_para = doc.add_paragraph()
                        env_para.add_run("Environment: ").bold = True
                        env_text = ", ".join(proj['environment']) if isinstance(proj['environment'], list) else str(proj['environment'])
                        env_para.add_run(env_text).font.size = Pt(9)
                        env_para.paragraph_format.space_after = Pt(8)
    
    # --- Education ---
    if 'education' in data and data['education']:
        add_section_heading("Education")
        for edu in data['education']:
            p = doc.add_paragraph()
            parts = []
            if edu.get('degree'):
                parts.append(edu['degree'])
            if edu.get('institution'):
                parts.append(edu['institution'])
            if edu.get('location'):
                parts.append(edu['location'])
            if edu.get('year'):
                parts.append(f"({edu['year']})")
            p.add_run(" | ".join(parts)).font.size = Pt(9)
    
    # Add very small space between Education and Certifications
    if 'education' in data and data['education'] and 'certifications' in data and data['certifications']:
        spacing_para = doc.add_paragraph()
        spacing_para.paragraph_format.space_after = Pt(2)
    
    # --- Certifications ---
    if 'certifications' in data and data['certifications']:
        add_section_heading("Certifications")
        for cert in data['certifications']:
            if cert:  # Skip empty strings
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(cert).font.size = Pt(9)
    
    doc.save(output_path)
    return output_path
#==========================================================================================================


